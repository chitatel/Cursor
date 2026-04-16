from __future__ import annotations

import argparse
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DEFAULT_SOURCE = Path(
    "/Users/chitatelsnov/Downloads/Описание процесса выплаты Гарантийного удержания.docx"
)
DEFAULT_OUTPUT = Path("Doc/gu_payment_process_rag.docx")
DEFAULT_PREVIEW = Path("Doc/gu_payment_process_rag_preview.txt")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a RAG-friendly DOCX instruction for guarantee holdback payments."
    )
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--preview", type=Path, default=DEFAULT_PREVIEW)
    return parser.parse_args()


def extract_images_in_order(source: Path, target_dir: Path) -> list[Path]:
    doc = Document(str(source))
    target_dir.mkdir(parents=True, exist_ok=True)

    ext_map = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
        "image/x-wmf": ".wmf",
        "image/x-emf": ".emf",
    }

    ordered_paths: list[Path] = []
    counter = 0

    for para in doc.paragraphs:
        for run in para.runs:
            blips = run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in blips:
                embed_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if not embed_id:
                    continue
                rel = doc.part.rels[embed_id]
                content_type = rel.target_part.content_type
                ext = ext_map.get(content_type, ".png")
                counter += 1
                image_path = target_dir / f"source_{counter:03d}{ext}"
                image_path.write_bytes(rel.target_part.blob)
                ordered_paths.append(image_path)

    return ordered_paths


def build_blocks(images: list[Path]) -> list[dict]:
    if len(images) != 11:
        raise ValueError(f"Expected 11 images in the source document, got {len(images)}")

    return [
        {
            "kind": "title",
            "text": "Инструкция по выплате гарантийного удержания (ГУ) для загрузки в RAG",
        },
        {
            "kind": "paragraph",
            "text": (
                "Назначение инструкции. Документ описывает полный порядок выплаты "
                "гарантийного удержания, действия в 1С:УХ и 1С:ДО, а также типовые "
                "ошибки, которые возникают при согласовании заявки на оплату."
            ),
        },
        {
            "kind": "paragraph",
            "text": (
                "Общая схема процесса выплаты гарантийного удержания: пользователь "
                "формирует ЗНО ГУ в 1С:УХ, система проверяет срок ГУ и сумму выплаты, "
                "при необходимости пользователь устраняет замечания, затем оформляет "
                "служебную записку по ГУ в 1С:ДО и после полного согласования повторно "
                "запускает ЗНО ГУ."
            ),
            "images": [images[0], images[1]],
        },
        {"kind": "heading", "text": "1. Проверка заявки на оплату в 1С:УХ"},
        {
            "kind": "paragraph",
            "text": (
                "1.1. Откройте или сформируйте ЗНО ГУ в 1С:УХ и отправьте документ "
                "на согласование. При запуске согласования система автоматически "
                "проверяет два условия: заполнена ли Дата РВ по объекту строительства "
                "и не превышена ли сумма выплаты по гарантийному удержанию."
            ),
        },
        {
            "kind": "paragraph",
            "text": (
                "1.2. Если при отправке появляется сообщение "
                "«Не найдены Даты РВ по проекту: ...», это означает, что в карточке "
                "объекта строительства не заполнено поле «Дата РВ». Пока дата РВ "
                "не будет заполнена, ЗНО ГУ не пройдет проверку и не уйдет "
                "на согласование."
            ),
            "images": [images[2]],
        },
        {
            "kind": "paragraph",
            "text": (
                "1.3. Чтобы исправить ошибку по Дате РВ, в 1С:УХ перейдите по пути: "
                "Справочники -> Управление НСИ -> Объекты строительства. В списке "
                "найдите объект строительства, который указан в сообщении об ошибке, "
                "и откройте его карточку."
            ),
            "images": [images[3]],
        },
        {
            "kind": "paragraph",
            "text": (
                "1.4. В карточке объекта строительства заполните поле «Дата РВ», "
                "сохраните изменения и закройте карточку. После этого вернитесь "
                "в ЗНО ГУ и повторно отправьте заявку на оплату на согласование."
            ),
            "images": [images[4]],
        },
        {
            "kind": "paragraph",
            "text": (
                "1.5. Если при отправке появляется сообщение "
                "«Превышен лимит гарантийных удержаний по договору. Согласованный "
                "лимит: <Сумма>», это означает, что система не позволяет выплатить "
                "текущую сумму без дополнительного согласования. В этом случае "
                "необходимо оформить служебную записку по ГУ в 1С:ДО."
            ),
            "images": [images[5]],
        },
        {"kind": "heading", "text": "2. Оформление служебной записки по ГУ в 1С:ДО"},
        {
            "kind": "paragraph",
            "text": (
                "2.1. В 1С:ДО откройте раздел «Документы и файлы» и создайте новый "
                "внутренний документ по шаблону «Служебная записка ГУ» по пути: "
                "Документы и файлы -> Создать -> Служебная записка ГУ -> Создать."
            ),
            "images": [images[6]],
        },
        {
            "kind": "paragraph",
            "text": (
                "2.2. В карточке внутреннего документа заполните обязательные поля. "
                "Для процесса выплаты гарантийного удержания обязательно проверьте "
                "реквизиты «Договор» и «Сумма ГУ к выплате». После заполнения "
                "нажмите кнопку «Зарегистрировать»."
            ),
            "images": [images[7]],
        },
        {
            "kind": "paragraph",
            "text": (
                "2.3. После регистрации 1С:ДО предложит сразу запустить процесс "
                "согласования. Нажмите «Перейти к запуску процесса», чтобы не "
                "возвращаться к документу вручную и сразу открыть карточку запуска "
                "маршрута."
            ),
            "images": [images[8]],
        },
        {
            "kind": "paragraph",
            "text": (
                "2.4. Если окно автоматического запуска не появилось или документ "
                "нужно отправить вручную, откройте зарегистрированный документ, "
                "нажмите «Отправить», выберите процесс «Служебная записка (ГУ)» "
                "и нажмите «Создать процесс»."
            ),
            "images": [images[9]],
        },
        {
            "kind": "paragraph",
            "text": (
                "2.5. В карточке комплексного процесса проверьте, что выбран нужный "
                "маршрут согласования, и нажмите кнопку «Стартовать и закрыть». "
                "После этого служебная записка по ГУ будет отправлена на согласование."
            ),
            "images": [images[10]],
        },
        {"kind": "heading", "text": "3. Повторный запуск ЗНО ГУ после согласования"},
        {
            "kind": "paragraph",
            "text": (
                "3.1. Дождитесь полного согласования служебной записки по ГУ в 1С:ДО. "
                "Документ должен пройти весь маршрут согласования без замечаний."
            ),
        },
        {
            "kind": "paragraph",
            "text": (
                "3.2. После полного согласования вернитесь в 1С:УХ, откройте исходное "
                "ЗНО ГУ и повторно отправьте заявку на оплату гарантийного удержания "
                "на согласование."
            ),
        },
        {"kind": "heading", "text": "4. Быстрые ответы по типовым ошибкам"},
        {
            "kind": "paragraph",
            "text": (
                "4.1. Ошибка «Не найдены Даты РВ по проекту» означает, что по объекту "
                "строительства не заполнено поле «Дата РВ». Нужно открыть карточку "
                "объекта строительства в 1С:УХ, заполнить «Дата РВ» и повторно "
                "отправить ЗНО ГУ."
            ),
        },
        {
            "kind": "paragraph",
            "text": (
                "4.2. Ошибка «Превышен лимит гарантийных удержаний по договору» "
                "означает, что нужно создать служебную записку по ГУ в 1С:ДО, "
                "заполнить договор и сумму ГУ к выплате, зарегистрировать документ, "
                "запустить процесс согласования и только после согласования повторно "
                "отправить ЗНО ГУ в 1С:УХ."
            ),
        },
    ]


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)


def add_block(doc: Document, block: dict) -> None:
    if block["kind"] == "title":
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(block["text"])
        run.bold = True
        run.font.size = Pt(14)
        return

    if block["kind"] == "heading":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block["text"])
        run.bold = True
        run.font.size = Pt(12)
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(block["text"])

    for image_path in block.get("images", []):
        paragraph.add_run().add_break()
        picture_run = paragraph.add_run()
        picture_run.add_picture(str(image_path), width=Inches(5.9))


def build_preview(blocks: list[dict]) -> str:
    lines: list[str] = []
    marker_counter = 0

    for block in blocks:
        lines.append(block["text"])
        for image_path in block.get("images", []):
            marker_counter += 1
            lines.append(f"[Рисунок {marker_counter}: img_{marker_counter:03d}{image_path.suffix}]")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def main() -> None:
    args = parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.preview.parent.mkdir(parents=True, exist_ok=True)

    with TemporaryDirectory(prefix="gu_rag_images_") as tmp_dir_name:
        image_dir = Path(tmp_dir_name)
        images = extract_images_in_order(args.source, image_dir)
        blocks = build_blocks(images)

        doc = Document()
        style_document(doc)
        for block in blocks:
            add_block(doc, block)
        doc.save(str(args.output))

        args.preview.write_text(build_preview(blocks), encoding="utf-8")


if __name__ == "__main__":
    main()
