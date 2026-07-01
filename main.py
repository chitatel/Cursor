"""
RAG сервис v5.0 — production pipeline
- FAISS (IndexFlatIP) + LlamaIndex SentenceSplitter
- Полноценный hybrid retrieval: FAISS vector search + глобальный BM25
  по всему корпусу, объединение через RRF fusion
- BM25 с prefix-stemming (5-char prefix) для морфологических вариантов
  русского языка ("согласовать" ↔ "согласования")
- Document-aware context: соседние чанки того же документа подтягиваются
  автоматически, чтобы ответы по большим документам не строились на
  изолированных фрагментах
- Адаптивный chunking (cheatsheet 350/50 vs manual 800/120)
- Overlap-эвристика защиты от галлюцинаций (без доп. LLM-вызова)
- Батчевые эмбеддинги (×N быстрее индексирования)
- Thread-safe in-memory кэш записей; index.json с recovery при повреждении
- _faiss_index сбрасывается в None ДО записи JSON и rebuild — при сбое
  rebuild поиск вернёт пустой результат, а не данные от старого индекса
- index.json пишется раньше index.faiss; FAISS всегда перестраивается из JSON
  на старте — защита от рассинхронизации в том числе при совпадении ntotal
- Бэкапы битого JSON именуются с timestamp (index.json.bak.<ts>), не перезаписываются
- Path traversal защита на всех файловых эндпоинтах (Python 3.8+ совместимо)
- Потоковая загрузка файлов (без буферизации в памяти)
"""

import os
import re
import json
import math
import time
import shutil
import logging
import asyncio
from pathlib import Path
from typing import Optional
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from zoneinfo import ZoneInfo

import httpx
import numpy as np
import aiofiles
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, BackgroundTasks, Request
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

APP_DIR = Path(__file__).resolve().parent
CONFIG_PATH = APP_DIR / "config.json"


def _load_config() -> dict:
    if not CONFIG_PATH.exists():
        return {}
    try:
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception as e:
        raise RuntimeError(f"Failed to read config file '{CONFIG_PATH}': {e}") from e


CONFIG = _load_config()


def _config_value(name: str):
    if name not in CONFIG:
        raise RuntimeError(f"Missing required config key '{name}' in '{CONFIG_PATH}'")
    return CONFIG[name]


OLLAMA_BASE_URL = _config_value("OLLAMA_BASE_URL")
OLLAMA_LLM_MODEL = _config_value("OLLAMA_LLM_MODEL")
OLLAMA_EMBED_MODEL = _config_value("OLLAMA_EMBED_MODEL")
LLM_API_MODE = str(_config_value("LLM_API_MODE")).strip().lower()
STORAGE_DIR = Path(_config_value("STORAGE_DIR"))
if not STORAGE_DIR.is_absolute():
    STORAGE_DIR = (APP_DIR / STORAGE_DIR).resolve()
FILES_DIR = STORAGE_DIR / "files"
ASK_LOG_FILE = STORAGE_DIR / "ask.log.jsonl"
PREPARED_DIR = STORAGE_DIR / "prepared"
INDEX_FILE = STORAGE_DIR / "index.json"
FAISS_INDEX_FILE = STORAGE_DIR / "index.faiss"
TOP_K = int(_config_value("TOP_K"))
CHUNK_SIZE = int(_config_value("CHUNK_SIZE"))
CHUNK_OVERLAP = int(_config_value("CHUNK_OVERLAP"))
BASE_URL = CONFIG.get("BASE_URL", "")


def _public_base_url(request: Request) -> str:
    """Берём BASE_URL из конфига, а если пустой — из входящего запроса."""
    if BASE_URL:
        return BASE_URL.rstrip("/")
    return str(request.base_url).rstrip("/")
SIM_THRESHOLD = float(_config_value("SIM_THRESHOLD"))
OLLAMA_API_KEY = _config_value("OLLAMA_API_KEY")
OPENWEBUI_AUTH_PATH = _config_value("OPENWEBUI_AUTH_PATH")
OPENWEBUI_USER = _config_value("OPENWEBUI_USER")
OPENWEBUI_PASSWORD = _config_value("OPENWEBUI_PASSWORD")

# Корпоративная База знаний (HTTP Basic Auth, Apache mod_auth)
KB_PORTAL_BASE_URL = str(CONFIG.get("KB_PORTAL_BASE_URL", "")).rstrip("/")
KB_PORTAL_USER = str(CONFIG.get("KB_PORTAL_USER", ""))
KB_PORTAL_PASSWORD = str(CONFIG.get("KB_PORTAL_PASSWORD", ""))
# Проверять SSL-сертификат портала?  Если у портала самоподписанный или
# подписанный внутренним корпоративным CA сертификат — поставить false.
KB_PORTAL_VERIFY_SSL = bool(CONFIG.get("KB_PORTAL_VERIFY_SSL", True))

# Отсканированные PDF (без текстового слоя) отвергаются на загрузке.
# Если общее число символов, извлечённых из PDF, меньше этого порога —
# считаем его сканом и возвращаем 400 с понятной ошибкой.
PDF_MIN_TEXT_CHARS = int(CONFIG.get("PDF_MIN_TEXT_CHARS", 200))

# Автоописание картинок через vision-LLM (Gemma 4 E2B и т.п.).
# Если включено — при индексации каждая картинка (PDF, docx, EMF) описывается
# LLM, и описание встраивается в текст рядом с маркером [Рисунок N: file].
# Это делает картинки поисковыми (по содержимому) и улучшает attach_images.
# Отключай, если модель не поддерживает vision или индексация слишком медленная.
IMAGE_DESCRIPTION_ENABLED = bool(CONFIG.get("IMAGE_DESCRIPTION_ENABLED", True))
# Пустая строка = использовать OLLAMA_LLM_MODEL.
IMAGE_DESCRIPTION_MODEL = str(CONFIG.get("IMAGE_DESCRIPTION_MODEL", "")).strip()

# Размер батча для эмбеддингов при индексировании
EMBED_BATCH_SIZE = int(CONFIG.get("EMBED_BATCH_SIZE", 32))
# Максимум записей, которое отдаёт GET /logs за один запрос
LOG_MAX_RETURN = int(CONFIG.get("LOG_MAX_RETURN", 5000))
# Включить query rewrite (добавляет 1 LLM-вызов, может улучшить recall)
QUERY_REWRITE_ENABLED = bool(CONFIG.get("QUERY_REWRITE_ENABLED", False))
# Вес BM25 при RRF fusion (0.0 — только вектор, 1.0 — только BM25)
# BM25 работает глобально по всему корпусу с prefix-stemming
BM25_WEIGHT = float(CONFIG.get("BM25_WEIGHT", 0.5))

os.environ["ANONYMIZED_TELEMETRY"] = "FALSE"

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

STORAGE_DIR.mkdir(exist_ok=True)
FILES_DIR.mkdir(exist_ok=True)
PREPARED_DIR.mkdir(exist_ok=True)

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".msg"}

# ── Глобальное состояние ─────────────────────────────────────────────────────

_records: Optional[list[dict]] = None
_records_lock = asyncio.Lock()
_faiss_index = None
_indexing: dict[str, dict] = {}
_indexing_locks: dict[str, asyncio.Lock] = {}
_auth_token: Optional[str] = None
_auth_lock = asyncio.Lock()
_ask_log_lock = asyncio.Lock()
LOCAL_TZ = ZoneInfo(CONFIG.get("LOCAL_TIMEZONE", "Europe/Volgograd"))


# ── Хранилище записей ────────────────────────────────────────────────────────

def _load_records_sync() -> list[dict]:
    """
    Синхронное чтение с диска (вызывать только под _records_lock).
    При повреждённом JSON переименовывает файл в index.json.bak.<timestamp>
    и стартует с пустым индексом, чтобы сервис поднялся, а не упал.
    Timestamp в имени гарантирует, что предыдущие бэкапы не перезаписываются.
    """
    global _records
    if _records is None:
        if INDEX_FILE.exists():
            try:
                _records = json.loads(INDEX_FILE.read_text(encoding="utf-8"))
            except Exception as e:
                ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                bak = INDEX_FILE.with_name(f"index.json.bak.{ts}")
                log.error(
                    "index.json is corrupted (%s). Renaming to %s and starting with empty index.",
                    e,
                    bak,
                )
                try:
                    INDEX_FILE.rename(bak)
                except Exception as rename_err:
                    log.error("Could not rename corrupted index.json: %s", rename_err)
                _records = []
        else:
            _records = []
    return _records


async def _get_records() -> list[dict]:
    """Потокобезопасное получение записей."""
    async with _records_lock:
        return list(_load_records_sync())


async def _log_ask(entry: dict) -> None:
    """
    Дописывает строку в JSON Lines лог /ask.  Запись атомарна на уровне строки
    благодаря O_APPEND; lock защищает от перемешивания с одновременных корутин
    в рамках одного воркера.
    """
    line = json.dumps(entry, ensure_ascii=False) + "\n"
    try:
        async with _ask_log_lock:
            async with aiofiles.open(ASK_LOG_FILE, "a", encoding="utf-8") as f:
                await f.write(line)
    except Exception as e:
        log.warning("Failed to write ask log: %s", e)


def _utc_iso_to_local_iso(value: str) -> str:
    """Конвертирует timestamp из UTC в локальную зону для отображения в 1С."""
    if not value:
        return value
    try:
        normalized = value.replace("Z", "+00:00")
        dt = datetime.fromisoformat(normalized)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(LOCAL_TZ).isoformat()
    except ValueError:
        return value


async def _read_ask_log(
    limit: int,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
) -> list[dict]:
    """
    Читает JSONL-лог и возвращает последние `limit` записей, новые сверху.
    date_from/date_to сравниваются как ISO-строки (лексикографически — это
    корректно для ISO-timestamps).
    """
    if not ASK_LOG_FILE.exists():
        return []
    entries: list[dict] = []
    async with _ask_log_lock:
        async with aiofiles.open(ASK_LOG_FILE, "r", encoding="utf-8") as f:
            async for raw in f:
                raw = raw.strip()
                if not raw:
                    continue
                try:
                    item = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                ts = item.get("timestamp", "")
                if date_from and ts < date_from:
                    continue
                if date_to and ts > date_to:
                    continue
                item["timestamp_utc"] = ts
                item["timestamp"] = _utc_iso_to_local_iso(ts)
                entries.append(item)
    entries.reverse()
    return entries[:limit]


def _save_records_sync(records: list[dict]):
    """Атомарная запись на диск (вызывать только под _records_lock)."""
    global _records
    tmp = INDEX_FILE.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(records, ensure_ascii=False), encoding="utf-8")
    tmp.replace(INDEX_FILE)
    _records = records


def _normalize_embedding(embedding: list[float]) -> list[float]:
    vec = np.asarray(embedding, dtype="float32")
    norm = np.linalg.norm(vec)
    if norm > 0:
        vec = vec / norm
    return vec.tolist()


# ── FAISS ────────────────────────────────────────────────────────────────────

def _load_faiss():
    import faiss

    return faiss


def _rebuild_faiss_index_sync(records: list[dict]):
    """Перестроить FAISS индекс (вызывать только под _records_lock)."""
    global _faiss_index
    if not records:
        _faiss_index = None
        if FAISS_INDEX_FILE.exists():
            FAISS_INDEX_FILE.unlink()
        return

    faiss = _load_faiss()
    dim = len(records[0]["embedding"])
    index = faiss.IndexFlatIP(dim)
    embeddings = np.asarray([r["embedding"] for r in records], dtype="float32")
    index.add(embeddings)
    tmp = FAISS_INDEX_FILE.with_suffix(".faiss.tmp")
    faiss.write_index(index, str(tmp))
    tmp.replace(FAISS_INDEX_FILE)
    _faiss_index = index


def _get_faiss_index_sync(records: list[dict]):
    """Получить FAISS индекс (вызывать только под _records_lock)."""
    global _faiss_index
    if not records:
        _faiss_index = None
        return None
    if _faiss_index is None:
        faiss = _load_faiss()
        if FAISS_INDEX_FILE.exists():
            index = faiss.read_index(str(FAISS_INDEX_FILE))
            if index.ntotal != len(records):
                log.warning(
                    "FAISS index mismatch: ntotal=%s, records=%s. Rebuilding.",
                    index.ntotal,
                    len(records),
                )
                _rebuild_faiss_index_sync(records)
            else:
                first = np.asarray(records[0]["embedding"], dtype="float32")
                norm = float(np.linalg.norm(first))
                if abs(norm - 1.0) > 0.01:
                    log.warning("Embeddings not normalized (norm=%.3f). Rebuilding.", norm)
                    for r in records:
                        r["embedding"] = _normalize_embedding(r["embedding"])
                    _save_records_sync(records)
                    _rebuild_faiss_index_sync(records)
                else:
                    _faiss_index = index
        else:
            _rebuild_faiss_index_sync(records)
    return _faiss_index


# ── Операции над записями (потокобезопасные) ─────────────────────────────────

async def _chunk_count() -> int:
    return len(await _get_records())


async def _document_chunk_counts() -> dict[str, int]:
    counts: dict[str, int] = {}
    for record in await _get_records():
        fn = record["filename"]
        counts[fn] = counts.get(fn, 0) + 1
    return counts


def _normalize_category(category: Optional[str]) -> Optional[str]:
    if category is None:
        return None
    cat = category.strip()
    return cat or None


async def _document_categories() -> dict[str, Optional[str]]:
    """Карта filename → category (берётся из первой встретившейся записи)."""
    result: dict[str, Optional[str]] = {}
    for record in await _get_records():
        fn = record["filename"]
        if fn not in result:
            result[fn] = record.get("category")
    return result


async def _list_categories() -> list[str]:
    """Уникальные категории, отсортированные."""
    cats: set[str] = set()
    for record in await _get_records():
        cat = record.get("category")
        if cat:
            cats.add(cat)
    return sorted(cats)


async def _document_category(filename: str) -> Optional[str]:
    for record in await _get_records():
        if record["filename"] == filename:
            return record.get("category")
    return None


async def _update_document_category(filename: str, category: Optional[str]) -> int:
    """Обновляет категорию у всех чанков документа без переиндексации."""
    async with _records_lock:
        records = _load_records_sync()
        updated = 0
        for r in records:
            if r["filename"] == filename:
                r["category"] = category
                updated += 1
        if updated:
            _save_records_sync(records)
        return updated


async def _document_records(filename: str) -> list[dict]:
    return [r for r in await _get_records() if r["filename"] == filename]


async def _replace_document_records(filename: str, new_records: list[dict]) -> int:
    global _faiss_index
    async with _records_lock:
        current = _load_records_sync()
        records = [r for r in current if r["filename"] != filename]
        removed = len(current) - len(records)
        records.extend(new_records)
        _faiss_index = None
        _save_records_sync(records)
        _rebuild_faiss_index_sync(records)
        return removed


async def _delete_document_records(filename: str) -> int:
    global _faiss_index
    async with _records_lock:
        current = _load_records_sync()
        records = [r for r in current if r["filename"] != filename]
        removed = len(current) - len(records)
        if removed == 0:
            return 0
        _faiss_index = None
        _save_records_sync(records)
        _rebuild_faiss_index_sync(records)
        return removed


async def _delete_category_records(category: str) -> dict[str, int]:
    """Удаляет из индекса все документы указанной категории за один rebuild."""
    global _faiss_index
    async with _records_lock:
        current = _load_records_sync()
        removed_by_file: dict[str, int] = {}
        records = []
        for record in current:
            if record.get("category") == category:
                fn = record["filename"]
                removed_by_file[fn] = removed_by_file.get(fn, 0) + 1
            else:
                records.append(record)
        if not removed_by_file:
            return {}
        _faiss_index = None
        _save_records_sync(records)
        _rebuild_faiss_index_sync(records)
        return removed_by_file


# ── Безопасность файловых путей ───────────────────────────────────────────────

def _safe_filename(filename: str) -> str:
    """
    Проверяет, что имя файла не выходит за пределы FILES_DIR.
    Возбуждает HTTPException(400) при path traversal попытке.
    Возвращает нормализованное имя файла (только basename).
    Совместимо с Python 3.8+ (не использует Path.is_relative_to).
    """
    name = Path(filename).name
    if not name or name != filename:
        raise HTTPException(400, f"Invalid filename: '{filename}'")
    resolved = str((FILES_DIR / name).resolve())
    files_dir = str(FILES_DIR.resolve())
    if not resolved.startswith(files_dir + os.sep) and resolved != files_dir:
        raise HTTPException(400, f"Invalid filename: '{filename}'")
    return name


def _bm25_scores(query: str, texts: list[str], k1: float = 1.5, b: float = 0.75) -> list[float]:
    """
    BM25 с prefix-stemming: для токенов ≥5 символов при отсутствии точного
    совпадения проверяется совпадение первых 5 символов.  Это покрывает
    основные морфологические вариации русского языка без внешнего стеммера.
    """
    _STEM_LEN = 5
    tokenize = lambda t: re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", t.lower())
    q_tokens = set(tokenize(query))
    if not q_tokens:
        return [0.0] * len(texts)

    tokenized = [tokenize(t) for t in texts]

    # Precompute per-document: exact token set + prefix set
    token_sets = [set(toks) for toks in tokenized]
    prefix_sets = [
        set(tok[:_STEM_LEN] for tok in toks if len(tok) >= _STEM_LEN)
        for toks in tokenized
    ]

    # Precompute DF for each query term (exact OR prefix match)
    df_cache: dict[str, int] = {}
    for term in q_tokens:
        df = 0
        term_prefix = term[:_STEM_LEN] if len(term) >= _STEM_LEN else None
        for i in range(len(tokenized)):
            if term in token_sets[i]:
                df += 1
            elif term_prefix and term_prefix in prefix_sets[i]:
                df += 1
        df_cache[term] = df

    doc_lens = [len(t) for t in tokenized]
    avg_dl = sum(doc_lens) / max(len(doc_lens), 1)
    n_docs = len(texts)

    scores = []
    for tokens in tokenized:
        tf_map: dict[str, int] = {}
        for tok in tokens:
            tf_map[tok] = tf_map.get(tok, 0) + 1
        prefix_tf: dict[str, int] = {}
        for tok in tokens:
            if len(tok) >= _STEM_LEN:
                p = tok[:_STEM_LEN]
                prefix_tf[p] = prefix_tf.get(p, 0) + 1

        score = 0.0
        dl = len(tokens)
        for term in q_tokens:
            df = df_cache[term]
            if df == 0:
                continue
            idf = math.log((n_docs - df + 0.5) / (df + 0.5) + 1)
            tf = tf_map.get(term, 0)
            if tf == 0 and len(term) >= _STEM_LEN:
                tf = prefix_tf.get(term[:_STEM_LEN], 0)
            tf_norm = tf * (k1 + 1) / (tf + k1 * (1 - b + b * dl / avg_dl))
            score += idf * tf_norm
        scores.append(score)
    return scores


def _filename_match_boost(query: str, filename: str) -> float:
    """
    Небольшой лексический boost по имени файла.
    Это помогает общим запросам вроде "как согласовать документ" поднимать
    файлы, в названии которых есть близкие ключи, не ломая основной ranking.

    Используется prefix-stemming (5 символов) — те же правила, что в BM25:
    "согласует" матчится с "согласующего" в имени файла.
    """
    _STEM_LEN = 5
    query_tokens = set(re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", query.lower()))
    if not query_tokens:
        return 0.0
    filename_tokens = set(re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", Path(filename).stem.lower()))
    if not filename_tokens:
        return 0.0

    exact_overlap = len(query_tokens & filename_tokens)

    # Prefix-overlap: совпадения по первым 5 буквам токенов длиной ≥5
    query_prefixes = {t[:_STEM_LEN] for t in query_tokens if len(t) >= _STEM_LEN}
    file_prefixes = {t[:_STEM_LEN] for t in filename_tokens if len(t) >= _STEM_LEN}
    prefix_overlap = len(query_prefixes & file_prefixes)

    overlap = max(exact_overlap, prefix_overlap)
    # Boost небольшой и дискретный: он только помогает shortlist, а не заменяет FAISS.
    return min(overlap * 0.35, 1.4)


def _rrf_fusion(
    vector_ids: list[str],
    bm25_ids: list[str],
    bm25_weight: float = 0.5,
    rrf_k: int = 60,
) -> list[str]:
    """
    Reciprocal Rank Fusion двух ранжированных списков.
    Возвращает список id, отсортированных по убыванию RRF-score.
    """
    scores: dict[str, float] = {}
    vector_weight = 1.0 - bm25_weight

    for rank, doc_id in enumerate(vector_ids):
        scores[doc_id] = scores.get(doc_id, 0.0) + vector_weight / (rrf_k + rank + 1)
    for rank, doc_id in enumerate(bm25_ids):
        scores[doc_id] = scores.get(doc_id, 0.0) + bm25_weight / (rrf_k + rank + 1)

    return sorted(scores, key=lambda x: scores[x], reverse=True)


async def _search_records(
    query: str,
    query_embedding: list[float],
    top_k: int,
    category: Optional[str] = None,
) -> dict:
    """Гибридный поиск: FAISS-кандидаты + глобальные lexical-кандидаты + RRF fusion.

    Если задана category — поиск ограничен только чанками этой категории.
    FAISS ищет по всему индексу, но результаты фильтруются по category;
    BM25 работает только по подмножеству, так что фильтр учитывается на обоих этапах.
    """
    async with _records_lock:
        records = _load_records_sync()
        index = _get_faiss_index_sync(records)

    if not records:
        return {
            "ids": [[]],
            "documents": [[]],
            "metadatas": [[]],
            "distances": [[]],
            "lexical_scores": [[]],
            "filename_scores": [[]],
        }

    if category is not None:
        allowed_ids = {r["id"] for r in records if r.get("category") == category}
        if not allowed_ids:
            return {
                "ids": [[]],
                "documents": [[]],
                "metadatas": [[]],
                "distances": [[]],
                "lexical_scores": [[]],
                "filename_scores": [[]],
            }
    else:
        allowed_ids = None

    id_to_record = {r["id"]: r for r in records}
    vector_ids: list[str] = []
    vector_distances: dict[str, float] = {}
    if index is not None:
        vec = np.asarray([query_embedding], dtype="float32")
        # При фильтре по категории берём больше кандидатов от FAISS, т.к. часть отсеется
        multiplier = 16 if allowed_ids is not None else 8
        vector_fetch_k = min(max(top_k * multiplier, 12), len(records))
        scores, indices = index.search(vec, vector_fetch_k)
        for score, idx in zip(scores[0], indices[0]):
            if idx < 0:
                continue
            record = records[int(idx)]
            if allowed_ids is not None and record["id"] not in allowed_ids:
                continue
            vector_ids.append(record["id"])
            vector_distances[record["id"]] = float(1 - score)

    # BM25 — только по подмножеству, если задан фильтр
    if allowed_ids is not None:
        lexical_records = [r for r in records if r["id"] in allowed_ids]
    else:
        lexical_records = records
    lexical_fetch_k = min(max(top_k * 8, 12), len(lexical_records))
    lexical_texts = [f"{record['filename']} {record['text']}" for record in lexical_records]
    bm25_scores = _bm25_scores(query, lexical_texts)
    filename_scores_by_id = {
        record["id"]: _filename_match_boost(query, record["filename"])
        for record in lexical_records
    }
    lexical_scores = [
        score + filename_scores_by_id.get(record["id"], 0.0)
        for score, record in zip(bm25_scores, lexical_records)
    ]
    lexical_scores_by_id = {
        record["id"]: float(score)
        for score, record in zip(lexical_scores, lexical_records)
    }
    lexical_ids = [
        lexical_records[i]["id"]
        for i in sorted(range(len(lexical_scores)), key=lambda x: lexical_scores[x], reverse=True)[:lexical_fetch_k]
        if lexical_scores[i] > 0
    ]

    fused_ids = _rrf_fusion(vector_ids, lexical_ids, bm25_weight=BM25_WEIGHT)[:top_k]

    ids: list[str] = []
    docs: list[str] = []
    metas: list[dict] = []
    distances: list[float] = []
    result_lexical_scores: list[float] = []
    result_filename_scores: list[float] = []
    for doc_id in fused_ids:
        if doc_id not in id_to_record:
            continue
        record = id_to_record[doc_id]
        ids.append(record["id"])
        docs.append(record["text"])
        metas.append({
            "filename": record["filename"],
            "chunk_index": record.get("chunk_index", 0),
        })
        distances.append(vector_distances.get(doc_id, 1.0))
        result_lexical_scores.append(lexical_scores_by_id.get(doc_id, 0.0))
        result_filename_scores.append(filename_scores_by_id.get(doc_id, 0.0))

    return {
        "ids": [ids],
        "documents": [docs],
        "metadatas": [metas],
        "distances": [distances],
        "lexical_scores": [result_lexical_scores],
        "filename_scores": [result_filename_scores],
    }


def _build_document_aware_context(
    result_ids: list[str],
    records: list[dict],
    budget: int,
) -> tuple[list[str], list[dict]]:
    """
    Расширяет найденные чанки соседними чанками того же документа (±1).
    Это критично для больших документов: если вопрос затрагивает процесс
    из нескольких шагов, ответ не должен строиться на изолированных фрагментах.

    Документы обрабатываются в порядке ранжирования — наиболее релевантный
    документ получает приоритет по бюджету.  Чанки внутри документа
    добавляются в порядке chunk_index (естественный порядок текста).

    budget — максимальное число чанков в итоговом контексте.
    """
    id_to_record = {r["id"]: r for r in records}

    # Сгруппировать выбранные чанки по документу (с сохранением порядка ранжирования)
    doc_selected: dict[str, list[int]] = {}
    doc_order: list[str] = []
    for rid in result_ids:
        rec = id_to_record.get(rid)
        if not rec:
            continue
        fn = rec["filename"]
        ci = rec.get("chunk_index", 0)
        if fn not in doc_selected:
            doc_selected[fn] = []
            doc_order.append(fn)
        doc_selected[fn].append(ci)

    # Построить карту чанков для каждого релевантного документа
    doc_chunks: dict[str, dict[int, dict]] = {}
    for rec in records:
        fn = rec["filename"]
        if fn not in doc_selected:
            continue
        ci = rec.get("chunk_index", 0)
        doc_chunks.setdefault(fn, {})[ci] = rec

    used = 0
    context_texts: list[str] = []
    context_metas: list[dict] = []

    for fn in doc_order:
        if used >= budget:
            break

        chunks_map = doc_chunks.get(fn, {})
        if not chunks_map:
            continue
        max_ci = max(chunks_map.keys())
        selected = sorted(doc_selected[fn])

        # Расширить окно: ±1 сосед для каждого выбранного чанка
        expanded = set()
        for ci in selected:
            for delta in range(-1, 2):
                neighbor = ci + delta
                if 0 <= neighbor <= max_ci and neighbor in chunks_map:
                    expanded.add(neighbor)

        # Добавить чанки в порядке их следования в документе
        for ci in sorted(expanded):
            if used >= budget:
                break
            rec = chunks_map[ci]
            context_texts.append(rec["text"])
            context_metas.append({"filename": fn, "chunk_index": ci})
            used += 1

    return context_texts, context_metas


# ── Auth / HTTP ───────────────────────────────────────────────────────────────

async def _get_auth_token(force_refresh: bool = False) -> str:
    global _auth_token

    if OLLAMA_API_KEY:
        return OLLAMA_API_KEY

    if not OPENWEBUI_PASSWORD or not OPENWEBUI_USER:
        return ""

    if _auth_token and not force_refresh:
        return _auth_token

    async with _auth_lock:
        if _auth_token and not force_refresh:
            return _auth_token

        signin_url = f"{_openwebui_root_url()}{OPENWEBUI_AUTH_PATH}"
        payload = {"user": OPENWEBUI_USER, "password": OPENWEBUI_PASSWORD}

        async with httpx.AsyncClient(timeout=60) as c:
            response = await c.post(
                signin_url,
                json=payload,
                headers={"Content-Type": "application/json"},
            )
            response.raise_for_status()
            data = response.json()

        token = (
            data.get("token")
            or data.get("access_token")
            or data.get("data", {}).get("token")
        )
        if not token and "token=" in response.headers.get("set-cookie", ""):
            cookie = response.headers["set-cookie"].split("token=", 1)[1]
            token = cookie.split(";", 1)[0]
        if not token:
            raise ValueError("Open WebUI signin succeeded but token was not returned")

        _auth_token = token
        return token


async def _api_headers(force_refresh: bool = False) -> dict:
    headers = {"Content-Type": "application/json"}
    token = await _get_auth_token(force_refresh=force_refresh)
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def _base_url() -> str:
    return OLLAMA_BASE_URL.rstrip("/")


def _api_mode() -> str:
    if LLM_API_MODE in {"ollama", "openai", "openwebui"}:
        return LLM_API_MODE
    base = _base_url().lower()
    if "/api/chat/completions" in base or "/ollama/api" in base:
        return "openwebui"
    if base.endswith("/v1") or "/api/v1" in base:
        return "openai"
    if OLLAMA_API_KEY:
        return "openwebui"
    return "ollama"


def _openai_base_url() -> str:
    base = _base_url()
    if base.endswith("/v1") or base.endswith("/api/v1"):
        return base
    return f"{base}/api/v1"


def _openwebui_root_url() -> str:
    base = _base_url()
    for suffix in ("/api/chat/completions", "/ollama/api/embed", "/ollama/api", "/api"):
        if base.endswith(suffix):
            return base[: -len(suffix)]
    return base


def _chat_url() -> str:
    mode = _api_mode()
    if mode == "openai":
        return f"{_openai_base_url()}/chat/completions"
    if mode == "openwebui":
        return f"{_openwebui_root_url()}/api/chat/completions"
    return f"{_base_url()}/api/chat"


def _openwebui_ollama_chat_url() -> str:
    return f"{_openwebui_root_url()}/ollama/api/chat"


def _embed_url() -> str:
    mode = _api_mode()
    if mode == "openai":
        return f"{_openai_base_url()}/embeddings"
    if mode == "openwebui":
        return f"{_openwebui_root_url()}/ollama/api/embed"
    return f"{_base_url()}/api/embed"


# ── Embed (с батчингом) ───────────────────────────────────────────────────────

async def _embed_one(text: str, client: httpx.AsyncClient) -> Optional[list[float]]:
    headers = await _api_headers()
    response = await client.post(
        _embed_url(),
        headers=headers,
        json={"model": OLLAMA_EMBED_MODEL, "input": text},
    )
    if response.status_code == 401 and OPENWEBUI_PASSWORD and OPENWEBUI_USER:
        headers = await _api_headers(force_refresh=True)
        response = await client.post(
            _embed_url(),
            headers=headers,
            json={"model": OLLAMA_EMBED_MODEL, "input": text},
        )
    if response.status_code == 400:
        return None
    response.raise_for_status()
    data = response.json()
    if "data" in data:
        return data["data"][0]["embedding"]
    embeddings = data["embeddings"]
    return embeddings[0] if isinstance(embeddings[0], list) else embeddings


async def _embed(text: str) -> Optional[list[float]]:
    """Получить эмбеддинг для одного текста."""
    async with httpx.AsyncClient(timeout=120) as c:
        return await _embed_one(text, c)


async def _embed_batch(texts: list[str]) -> list[Optional[list[float]]]:
    """
    Батчевые эмбеддинги: один HTTP-запрос на батч (если модель поддерживает),
    с fallback на поштучный запрос при ошибке.
    """
    if not texts:
        return []

    async with httpx.AsyncClient(timeout=300) as c:
        headers = await _api_headers()
        try:
            response = await c.post(
                _embed_url(),
                headers=headers,
                json={"model": OLLAMA_EMBED_MODEL, "input": texts},
            )
            if response.status_code == 401 and OPENWEBUI_PASSWORD and OPENWEBUI_USER:
                headers = await _api_headers(force_refresh=True)
                response = await c.post(
                    _embed_url(),
                    headers=headers,
                    json={"model": OLLAMA_EMBED_MODEL, "input": texts},
                )

            if response.status_code == 400:
                raise ValueError("batch not supported")

            response.raise_for_status()
            data = response.json()

            if "data" in data:
                sorted_items = sorted(data["data"], key=lambda x: x.get("index", 0))
                return [item["embedding"] for item in sorted_items]

            embeddings = data.get("embeddings", [])
            if embeddings and isinstance(embeddings[0], list):
                return embeddings
            raise ValueError("unexpected batch response shape")

        except Exception as e:
            log.warning("Batch embed failed (%s), falling back to sequential", e)
            results = []
            for text in texts:
                results.append(await _embed_one(text, c))
            return results


# ── LLM chat ─────────────────────────────────────────────────────────────────

async def _chat(
    messages: list[dict],
    max_tokens: int = 400,
    temperature: float = 0.0,
) -> str:
    async with httpx.AsyncClient(timeout=600) as c:
        headers = await _api_headers()
        mode = _api_mode()

        payload_openai = {
            "model": OLLAMA_LLM_MODEL,
            "messages": messages,
            "stream": False,
            "temperature": temperature,
            "top_p": 0.85,
            "max_tokens": max_tokens,
        }
        # Для «живого» стиля (temperature >= 0.3) ослабляем top_p/top_k, чтобы
        # модель могла выбирать более разнообразные формулировки. Для
        # детерминированных задач (captioning, prepare) сохраняем узкую выборку.
        if temperature >= 0.3:
            top_p = 0.95
            top_k = 60
            min_p = 0.05
            repeat_penalty = 1.1
        else:
            top_p = 0.85
            top_k = 30
            min_p = 0.0
            repeat_penalty = 1.15
        payload_ollama = {
            "model": OLLAMA_LLM_MODEL,
            "messages": messages,
            "stream": False,
            "think": False,
            "options": {
                "temperature": temperature,
                "top_p": top_p,
                "top_k": top_k,
                "min_p": min_p,
                "repeat_penalty": repeat_penalty,
                "num_predict": max_tokens,
                "num_ctx": 8192,
            },
        }

        async def _post(url: str, payload: dict):
            response = await c.post(url, headers=headers, json=payload)
            if response.status_code == 401 and OPENWEBUI_PASSWORD and OPENWEBUI_USER:
                refreshed_headers = await _api_headers(force_refresh=True)
                response = await c.post(url, headers=refreshed_headers, json=payload)
            if response.status_code >= 400:
                log.error("Chat API error %s at %s: %s", response.status_code, url, response.text)
            response.raise_for_status()
            return response

        if mode == "openai":
            response = await _post(_chat_url(), payload_openai)
            content = response.json()["choices"][0]["message"]["content"]
        elif mode == "openwebui":
            response = await _post(_openwebui_ollama_chat_url(), payload_ollama)
            content = response.json().get("message", {}).get("content", "")
        else:
            response = await _post(_chat_url(), payload_ollama)
            content = response.json().get("message", {}).get("content", "")

        if not content:
            raise ValueError("LLM вернул пустой ответ")
        return content


# ── Image captioning (vision LLM) ────────────────────────────────────────────
#
# Каждая картинка при индексации проходит через vision-модель (Gemma 4 E2B
# и т.п.) — получаем краткое описание, которое встраиваем в текст документа.
# Кэш по SHA256 в STORAGE_DIR/image_captions.json, чтобы реиндексация не
# описывала уже виденные картинки повторно.

_caption_cache: Optional[dict[str, str]] = None
_caption_cache_dirty: bool = False


def _caption_cache_file() -> Path:
    return STORAGE_DIR / "image_captions.json"


def _load_caption_cache() -> dict[str, str]:
    global _caption_cache
    if _caption_cache is None:
        p = _caption_cache_file()
        try:
            if p.exists():
                _caption_cache = json.loads(p.read_text(encoding="utf-8"))
            else:
                _caption_cache = {}
        except Exception:
            _caption_cache = {}
    return _caption_cache


def _save_caption_cache() -> None:
    global _caption_cache_dirty
    if not _caption_cache_dirty or _caption_cache is None:
        return
    p = _caption_cache_file()
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(
            json.dumps(_caption_cache, ensure_ascii=False),
            encoding="utf-8",
        )
        _caption_cache_dirty = False
    except Exception as e:
        log.warning("[image_caption] cache save failed: %s", e)


_IMAGE_CAPTION_PROMPT = (
    "Опиши изображение на русском языке для поиска. "
    "Обязательно ПРОЦИТИРУЙ дословно весь заметный текст на "
    "изображении: заголовки, подписи, метки блоков, названия "
    "кнопок, значения — всё что можно прочитать. Цитируй "
    "русские слова точно, как на картинке. "
    "После цитат добавь одно короткое предложение с типом "
    "изображения (блок-схема / инфографика / скриншот / "
    "иллюстрация / логотип / диаграмма) и его назначением. "
    "Без вступительных фраз вроде «на изображении», сразу текст."
)


async def _describe_image_async(image_bytes: bytes, model: str) -> str:
    """Запрашивает описание картинки у vision-LLM через chat API."""
    import base64
    b64 = base64.b64encode(image_bytes).decode("ascii")

    async with httpx.AsyncClient(timeout=120) as c:
        headers = await _api_headers()
        mode = _api_mode()

        if mode == "openai":
            url = f"{_openai_base_url()}/chat/completions"
            payload = {
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{b64}",
                            },
                        },
                        {"type": "text", "text": _IMAGE_CAPTION_PROMPT},
                    ],
                }],
                "stream": False,
                "temperature": 0.0,
                "max_tokens": 400,
            }
        else:
            if mode == "openwebui":
                url = _openwebui_ollama_chat_url()
            else:
                url = _chat_url()
            # Ollama chat API: images ставим ПЕРЕД текстом (Gemma 4 doc).
            payload = {
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": _IMAGE_CAPTION_PROMPT,
                    "images": [b64],
                }],
                "stream": False,
                "think": False,
                "options": {
                    "temperature": 0.0,
                    "num_predict": 400,
                    "num_ctx": 4096,
                },
            }

        response = await c.post(url, headers=headers, json=payload)
        if response.status_code == 401 and OPENWEBUI_PASSWORD and OPENWEBUI_USER:
            headers = await _api_headers(force_refresh=True)
            response = await c.post(url, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        if "choices" in data:
            return data["choices"][0]["message"]["content"].strip()
        return data.get("message", {}).get("content", "").strip()


def _describe_image_sync(image_path: Path) -> str:
    """
    Синхронная обёртка (для использования из executor thread во время
    индексации). Кэширует описания по SHA256 картинки. Возвращает пустую
    строку при ошибке или если captioning отключён.
    """
    if not IMAGE_DESCRIPTION_ENABLED:
        return ""
    try:
        data = image_path.read_bytes()
    except Exception:
        return ""
    if not data:
        return ""

    import hashlib
    sha = hashlib.sha256(data).hexdigest()

    cache = _load_caption_cache()
    if sha in cache:
        return cache[sha]

    model = IMAGE_DESCRIPTION_MODEL or OLLAMA_LLM_MODEL
    try:
        # Функция может вызываться и из sync executor thread (индексация
        # PDF/docx), и из async функции (KB-статьи _fetch_kb_article).
        # asyncio.run() валится в running loop → в этом случае запускаем
        # корутину в отдельном thread'е со своим event loop.
        try:
            asyncio.get_running_loop()
            in_running_loop = True
        except RuntimeError:
            in_running_loop = False

        if in_running_loop:
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                future = pool.submit(
                    asyncio.run, _describe_image_async(data, model)
                )
                caption = future.result(timeout=180)
        else:
            caption = asyncio.run(_describe_image_async(data, model))
    except Exception as e:
        log.warning(
            "[image_caption] %s failed (model=%s): %s",
            image_path.name, model, e,
        )
        return ""

    caption = (caption or "").strip()
    if not caption:
        return ""

    # ограничим по длине на всякий случай (модель может выдать много).
    # 1200 — с учётом транскрипции текста на инфографике / блок-схемах.
    if len(caption) > 1200:
        caption = caption[:1200].rsplit(" ", 1)[0] + "…"

    cache[sha] = caption
    global _caption_cache_dirty
    _caption_cache_dirty = True
    log.info(
        "[image_caption] %s → %d chars (model=%s)",
        image_path.name, len(caption), model,
    )
    return caption


def _build_image_marker(index: int, filename: str, image_path: Path) -> str:
    """
    Формирует маркер картинки для встраивания в текст документа.
    Если vision captioning включён и удалось получить описание —
    добавляет строку 'Изображение: <описание>' ПЕРЕД маркером.
    Это критично для attach_images: он ищет группы по словам ПЕРЕД
    маркером в чанке. Описание перед маркером = слова описания
    попадают в preceding_bag того же рисунка, и bag-of-words matching
    точно привяжет картинку к пункту ответа со связанным смыслом.
    """
    marker = f"[Рисунок {index}: {filename}]"
    caption = _describe_image_sync(image_path)
    if caption:
        return f"Изображение: {caption}\n{marker}"
    return marker


# ── Query rewrite ─────────────────────────────────────────────────────────────

async def _rewrite_query(q: str) -> str:
    """
    Query rewrite отключён.  Морфологические вариации покрываются
    prefix-stemming в BM25, а семантические — FAISS.  Дополнительный
    LLM-вызов на rewrite добавляет latency без гарантированного выигрыша
    по recall и может подменять смысл вопроса.
    """
    return q


# ── Document preparation ──────────────────────────────────────────────────────

PREPARE_DOCUMENT_PROMPT = """Ты готовишь документ для RAG-индексации.

Правила:
1. Не добавляй новых фактов и не удаляй важные детали.
2. Один логический шаг = один абзац.
3. Длинные абзацы разбивай на блоки до 700 символов.
4. Пошаговые инструкции оформляй нумерованным списком.
5. Заголовки оставляй отдельными строками.
6. Маркеры вида [Рисунок N: имя] сохрани строго без изменений.
7. Перед каждым маркером [Рисунок N: имя] должен быть текст с 2-4 значимыми словами, описывающими действие, экран, кнопку или поле.
8. Если перед картинкой нет контекста, добавь короткую фразу из ближайшего текста.
9. Табличные инструкции преобразуй в обычные списки.
10. Не используй заголовки вида "Рисунок 1" или "Рис. 2".

Верни только подготовленный текст."""


def _strip_code_fence(text: str) -> str:
    stripped = text.strip()
    if not stripped.startswith("```"):
        return stripped
    stripped = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", stripped)
    stripped = re.sub(r"\s*```$", "", stripped)
    return stripped.strip()


def _split_text_for_prepare(text: str, max_chars: int = 6000) -> list[str]:
    paragraphs = re.split(r"\n{2,}", text.strip())
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if current and current_len + len(paragraph) + 2 > max_chars:
            chunks.append("\n\n".join(current))
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += len(paragraph) + 2

    if current:
        chunks.append("\n\n".join(current))
    return chunks or [text.strip()]


async def _normalize_text_for_rag(text: str) -> str:
    parts = _split_text_for_prepare(text)
    normalized: list[str] = []
    for idx, part in enumerate(parts, start=1):
        content = await _chat(
            [
                {"role": "system", "content": PREPARE_DOCUMENT_PROMPT},
                {
                    "role": "user",
                    "content": (
                        f"Фрагмент {idx}/{len(parts)}. Подготовь его по правилам.\n\n"
                        f"{part}"
                    ),
                },
            ],
            max_tokens=3500,
        )
        normalized.append(_strip_code_fence(content))
    return "\n\n".join(p for p in normalized if p).strip()


def _significant_words(text: str) -> list[str]:
    return re.findall(r"[A-Za-zА-Яа-яЁё0-9]{4,}", text.lower())


def _image_markers(text: str) -> list[str]:
    return re.findall(r"\[Рисунок\s+\d+:\s+[^\]]+\]", text)


def _validate_prepared_text(text: str, source_filename: str, source_text: str = "") -> list[str]:
    warnings: list[str] = []
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]

    for idx, paragraph in enumerate(paragraphs, start=1):
        if len(paragraph) > 800:
            warnings.append(f"Абзац {idx} длиннее 800 символов ({len(paragraph)}).")

    marker_re = re.compile(r"\[Рисунок\s+(\d+):\s+([^\]]+)\]")
    markers = list(marker_re.finditer(text))
    if source_text:
        source_markers = set(_image_markers(source_text))
        prepared_markers = set(_image_markers(text))
        missing = sorted(source_markers - prepared_markers)
        extra = sorted(prepared_markers - source_markers)
        if missing:
            warnings.append(f"LLM потеряла маркеры изображений: {', '.join(missing[:5])}.")
        if extra:
            warnings.append(f"LLM добавила новые маркеры изображений: {', '.join(extra[:5])}.")

    seen_markers: set[str] = set()
    for marker in markers:
        marker_text = marker.group(0)
        if marker_text in seen_markers:
            warnings.append(f"Повторяется маркер изображения: {marker_text}.")
        seen_markers.add(marker_text)

        before = text[max(0, marker.start() - 220): marker.start()]
        words = _significant_words(before)
        if len(words) < 2:
            warnings.append(f"Перед {marker_text} меньше 2 значимых слов контекста.")

    bad_caption_re = re.compile(r"(?im)^\s*(рисунок|рис\.)\s*\d+\.?\s*$")
    for match in bad_caption_re.finditer(text):
        warnings.append(f"Найден отдельный заголовок '{match.group(0).strip()}', он может ломать привязку картинок.")

    stem_words = _significant_words(Path(source_filename).stem)
    if len(stem_words) < 2:
        warnings.append("Имя файла слишком короткое или не отражает тему документа.")

    return warnings


def _prepared_filename(filename: str) -> str:
    source = Path(filename)
    stem = re.sub(r"[^A-Za-zА-Яа-яЁё0-9_. -]+", "_", source.stem).strip(" ._")
    if not stem:
        stem = "document"
    return f"{stem}_prepared.docx"


def _add_text_to_docx(doc, text: str) -> None:
    for block in re.split(r"\n{2,}", text.strip()):
        block = block.strip()
        if block:
            doc.add_paragraph(block)


def _write_prepared_docx(text: str, out_path: Path, images_dir: Path) -> None:
    from docx import Document as DocxDocument
    from docx.shared import Inches

    doc = DocxDocument()
    marker_re = re.compile(r"\[Рисунок\s+\d+:\s+([^\]]+)\]")
    pos = 0

    for match in marker_re.finditer(text):
        _add_text_to_docx(doc, text[pos: match.start()])
        img_name = Path(match.group(1)).name
        img_path = images_dir / img_name
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(6.0))
            except Exception:
                doc.add_paragraph(match.group(0))
        else:
            doc.add_paragraph(match.group(0))
        pos = match.end()

    _add_text_to_docx(doc, text[pos:])
    doc.save(str(out_path))


# ── Faithfulness check ────────────────────────────────────────────────────────

def _is_faithful(answer: str, context: str) -> bool:
    """
    Быстрая лексическая проверка: ответ должен опираться на лексику контекста.
    Работает без LLM-вызова — не увеличивает латентность.
    Контекст не обрезается: проверяем по всему тексту, а не первым 3000 символам.
    """

    def normalize(text: str) -> str:
        return re.sub(r"\s+", " ", text.lower()).strip()

    ctx_norm = normalize(context)
    ans_norm = normalize(answer)
    words = [w for w in re.findall(r"[A-Za-zА-Яа-яЁё0-9-]{4,}", ans_norm)]
    if not words:
        return True
    overlap = sum(1 for w in words if w in ctx_norm)
    ratio = overlap / len(words)
    if ratio < 0.3:
        log.warning(
            "Faithfulness check: low overlap %.0f%% (%d/%d words)",
            ratio * 100,
            overlap,
            len(words),
        )
        return False
    return True


def _query_terms(text: str) -> set[str]:
    stop_words = {
        "какой", "какая", "какие", "какое", "кто", "что", "где", "куда", "когда",
        "если", "или", "для", "при", "про", "как", "это", "этот", "эта", "эти",
        "документ", "документа", "документы", "файл", "файла",
    }
    terms: set[str] = set()
    for word in re.findall(r"[A-Za-zА-Яа-яЁё0-9]{4,}", text.lower()):
        if word not in stop_words:
            terms.add(word)
            if len(word) >= 5:
                terms.add(word[:5])
    return terms


def _sentence_terms(text: str) -> set[str]:
    terms: set[str] = set()
    for word in re.findall(r"[A-Za-zА-Яа-яЁё0-9]{4,}", text.lower()):
        terms.add(word)
        if len(word) >= 5:
            terms.add(word[:5])
    return terms


def _is_broad_query(question: str) -> bool:
    """
    Широкий запрос = «название инструкции целиком», без конкретики.
    На такой запрос надо отдать полный контекст и разрешить объединять
    разделы из разных документов, иначе target_section_context выдернет
    одно предложение из середины.
    """
    q = (question or "").strip().lower()
    if not q:
        return False
    if "?" in q:
        return False
    # Триггеры конкретики — точно НЕ широкий запрос
    specific_markers = (
        "как ", "где ", "когда ", "почему ", "что делать", "что значит",
        "не работает", "не могу", "не получается", "ошибк", "проблем",
        "сообщение", "превышен", "не найден", "нельзя", "запрещ",
    )
    if any(marker in q for marker in specific_markers):
        return False
    if re.search(r"[«»\"']", q):
        return False
    # Считаем «содержательные» слова (длиннее 3 букв), без стоп-слов
    words = re.findall(r"[A-Za-zА-Яа-яЁё]{4,}", q)
    return 1 <= len(words) <= 4


def _most_relevant_fragments(
    query: str,
    docs: list[str],
    metas: list[dict],
    limit: int = 5,
) -> list[str]:
    q_terms = _query_terms(query)
    if not q_terms:
        return []

    scored: list[tuple[float, int, str]] = []
    for doc_idx, (doc_text, meta) in enumerate(zip(docs, metas)):
        sentences = [
            part.strip()
            for part in re.split(r"(?<=[.!?])\s+|\n+", doc_text)
            if part.strip()
        ]
        for sentence in sentences:
            s_terms = _sentence_terms(sentence)
            overlap = len(q_terms & s_terms)
            if overlap == 0:
                continue
            score = overlap + 1 / (doc_idx + 1)
            filename = meta.get("filename", "")
            chunk_index = meta.get("chunk_index", 0)
            scored.append((score, doc_idx, f"[{filename} | chunk {chunk_index}] {sentence}"))

    scored.sort(key=lambda item: (-item[0], item[1]))
    fragments: list[str] = []
    seen: set[str] = set()
    for _, _, fragment in scored:
        if fragment in seen:
            continue
        seen.add(fragment)
        fragments.append(fragment)
        if len(fragments) >= limit:
            break
    return fragments


def _looks_like_section_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 140:
        return False
    low = stripped.lower()
    if low.startswith(("как ", "что ", "где ", "когда ")):
        return True
    if re.match(r"^\d+[\.\)]\s+\S", stripped) and not stripped.endswith((".", ";", ":")):
        return True
    return False


def _target_section_context(
    query: str,
    docs: list[str],
    metas: list[dict],
    max_chars: int = 2500,
) -> str:
    q_terms = _query_terms(query)
    if not q_terms:
        return ""

    best: tuple[float, int, int] | None = None
    best_lines: list[str] = []
    best_meta: dict = {}

    for doc_idx, (doc_text, meta) in enumerate(zip(docs, metas)):
        lines = [line.strip() for line in doc_text.splitlines() if line.strip()]
        if not lines:
            continue
        for line_idx, line in enumerate(lines):
            line_terms = _sentence_terms(line)
            overlap = len(q_terms & line_terms)
            if overlap == 0:
                continue
            score = overlap
            if _looks_like_section_heading(line):
                score += 2
            if query.lower().strip() in line.lower():
                score += 3
            ranked = (float(score), -doc_idx, -line_idx)
            if best is None or ranked > best:
                best = ranked
                best_lines = lines
                best_meta = meta

    if best is None or not best_lines:
        return ""

    start_idx = -best[2]
    collected: list[str] = []
    for idx in range(start_idx, len(best_lines)):
        line = best_lines[idx]
        if idx > start_idx and collected and _looks_like_section_heading(line):
            break
        collected.append(line)
        if len("\n".join(collected)) >= max_chars:
            break

    if not collected:
        return ""

    filename = best_meta.get("filename", "")
    chunk_index = best_meta.get("chunk_index", 0)
    return f"[{filename} | chunk {chunk_index}]\n" + "\n".join(collected)


def _document_support_scores(
    answer: str,
    docs: list[str],
    metas: list[dict],
) -> dict[str, float]:
    answer_terms = _query_terms(answer)
    if not answer_terms:
        return {}

    scores: dict[str, float] = {}
    for doc_idx, (doc_text, meta) in enumerate(zip(docs, metas)):
        filename = meta.get("filename", "")
        doc_terms = _sentence_terms(doc_text)
        overlap = len(answer_terms & doc_terms)
        if overlap == 0:
            continue
        coverage = overlap / max(len(answer_terms), 1)
        score = overlap + coverage * 10 + 1 / (doc_idx + 1)
        if score > scores.get(filename, 0):
            scores[filename] = score
    return scores


# ── Форматирование ответа в HTML ─────────────────────────────────────────────

def _answer_to_html(
    answer: str,
    image_urls: dict[str, str],
    download_urls: dict[str, str] | None = None,
) -> str:
    """
    Конвертирует текстовый ответ в HTML для отображения в 1С.
    - Нумерованные пункты → <ol><li>
    - URL картинок → <img> теги
    - Ссылки на документы перед ответом
    - Остальной текст → <p>
    """
    from html import escape

    lines = answer.split("\n")
    html_parts: list[str] = []
    in_list = False
    numbered_re = re.compile(r"^(\d+)[\.\)]\s+(.*)")

    url_set = set(image_urls.values()) if image_urls else set()

    # Блок ссылок на документы-источники перед ответом
    if download_urls:
        source_items = list(download_urls.items())
        visible_sources = source_items[:2]
        hidden_sources = source_items[2:]

        html_parts.append('<p style="margin-top:0; font-size:13px; color:#666;">Вам может подойти:</p>')
        html_parts.append('<ul style="padding-left:20px; font-size:13px;">')
        for filename, url in visible_sources:
            html_parts.append(
                f'<li><a href="{escape(url)}" style="color:#1a73e8; text-decoration:none;">'
                f'{escape(filename)}</a></li>'
            )
        html_parts.append("</ul>")

        if hidden_sources:
            html_parts.append(
                '<details style="margin-top:-4px; margin-left:20px; font-size:13px;">'
                '<summary style="cursor:pointer; color:#1a73e8;">Еще</summary>'
            )
            html_parts.append('<ul style="padding-left:20px; margin-top:8px;">')
            for filename, url in hidden_sources:
                html_parts.append(
                    f'<li><a href="{escape(url)}" style="color:#1a73e8; text-decoration:none;">'
                    f'{escape(filename)}</a></li>'
                )
            html_parts.append("</ul></details>")

        html_parts.append('<hr style="margin:14px 0 16px; border:none; border-top:1px solid #ddd;">')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Строка — это URL картинки
        if stripped in url_set:
            if in_list:
                html_parts.append(
                    f'<br><img src="{escape(stripped)}" '
                    f'style="max-width:100%; margin-top:8px; border:1px solid #ccc; border-radius:4px;">'
                )
            else:
                html_parts.append(
                    f'<p><img src="{escape(stripped)}" '
                    f'style="max-width:100%; border:1px solid #ccc; border-radius:4px;"></p>'
                )
            continue

        m = numbered_re.match(stripped)
        if m:
            if not in_list:
                html_parts.append("<ol>")
                in_list = True
            html_parts.append(f"<li>{escape(m.group(2))}")
            continue

        # Обычная строка
        if in_list:
            html_parts.append("</li></ol>")
            in_list = False
        html_parts.append(f"<p>{escape(stripped)}</p>")

    if in_list:
        html_parts.append("</li></ol>")

    body = "\n".join(html_parts)
    return (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        "<style>"
        "body { font-family: Arial, sans-serif; font-size: 14px; padding: 12px; color: #333; }"
        "ol { padding-left: 20px; }"
        "li { margin-bottom: 10px; }"
        "p { margin: 8px 0; }"
        "img { display: block; }"
        "a:hover { text-decoration: underline !important; }"
        "</style>"
        f"</head><body>{body}</body></html>"
    )


# ── Привязка иллюстраций к пунктам ответа ────────────────────────────────────

def _attach_images_to_answer(
    answer: str,
    chunks: list[str],
    chunk_metas: list[dict],
    image_urls: dict[str, str],
) -> str:
    """
    Группируем картинки, идущие подряд в чанке (без значимого текста между ними).
    Для каждой группы берём текст перед первой картинкой группы.
    Для каждого пункта ответа ищем группу с наибольшим пересечением слов.
    Привязываем всю группу (все картинки) к пункту.

    Дополнительно, для случаев когда пункт ответа явно ссылается на имя
    файла-источника (например, «...представлена в Приложении А СТО.CNt-008.1.25...»),
    делаем второй проход и привязываем оставшиеся картинки по имени файла.
    """
    img_marker_re = re.compile(r"\[Рисунок \d+: [^\]]+\]")

    def bag(text: str) -> set[str]:
        cleaned = img_marker_re.sub("", text)
        return set(re.findall(r"[А-Яа-яЁёA-Za-z0-9-]{4,}", cleaned.lower()))

    def filename_tokens(filename: str) -> set[str]:
        """Осмысленные слова из имени файла: CTO.CNt-008.1.25, термолэнд, и т.п."""
        stem = filename.rsplit(".", 1)[0]
        tokens = set(re.findall(r"[А-Яа-яЁёA-Za-z0-9\.\-]{3,}", stem.lower()))
        return tokens

    # Группируем маркеры: если между двумя маркерами нет значимого текста
    # (менее 3 слов), они образуют одну группу
    # group = (preceding_bag, [url1, url2, ...], source_filename)
    groups: list[tuple[set[str], list[str], str]] = []
    # Сопоставление chunk_text → meta (для извлечения filename группы)
    text_to_filename: dict[int, str] = {}
    for idx, (chunk_text, meta) in enumerate(zip(chunks, chunk_metas)):
        text_to_filename[id(chunk_text)] = meta.get("filename", "")

    for chunk_text in chunks:
        markers_in_chunk = list(img_marker_re.finditer(chunk_text))
        if not markers_in_chunk:
            continue

        chunk_filename = text_to_filename.get(id(chunk_text), "")

        current_group_urls: list[str] = []
        current_group_preceding: set[str] = set()

        for i, match in enumerate(markers_in_chunk):
            url = image_urls.get(match.group(0))
            if not url:
                continue

            start = markers_in_chunk[i - 1].end() if i > 0 else 0
            preceding_text = chunk_text[start:match.start()]
            preceding_words = bag(preceding_text)

            if i == 0 or len(preceding_words) >= 1:
                # Начало новой группы: первый маркер или есть хоть одно слово перед ним
                if current_group_urls:
                    groups.append((current_group_preceding, current_group_urls, chunk_filename))
                current_group_preceding = preceding_words
                current_group_urls = [url]
            else:
                # Продолжение группы: между маркерами мало текста
                current_group_urls.append(url)

        if current_group_urls:
            groups.append((current_group_preceding, current_group_urls, chunk_filename))

    log.info(
        "[attach_images] %d groups built from %d chunks (image_urls: %d)",
        len(groups), len(chunks), len(image_urls),
    )

    if not groups:
        return answer

    # Разбиваем ответ на пункты
    point_re = re.compile(r"(?:^|\n)(\d+[\.\)]\s+.*?)(?=\n\d+[\.\)]\s+|\Z)", re.DOTALL)
    points = point_re.findall(answer)
    if not points:
        points = [line for line in answer.split("\n") if line.strip()]
    if not points:
        return answer

    log.info("[attach_images] %d points parsed from answer", len(points))

    used_groups: set[int] = set()
    result_parts: list[str] = []
    point_assignments: list[int] = []  # для каждого пункта — индекс группы (-1 если не нашли)

    # Первый проход: bag-of-words по тексту вокруг картинки (исходная логика)
    for pi, point in enumerate(points):
        point_bag = bag(point)
        if not point_bag:
            point_assignments.append(-1)
            continue

        best_idx = -1
        best_score = 0
        for idx, (preceding_bag, urls, _fn) in enumerate(groups):
            if idx in used_groups:
                continue
            score = len(point_bag & preceding_bag)
            if score > best_score:
                best_score = score
                best_idx = idx

        if best_idx >= 0 and best_score >= 2:
            used_groups.add(best_idx)
            point_assignments.append(best_idx)
            log.info(
                "[attach_images] point %d → group %d (bag-of-words score=%d)",
                pi + 1, best_idx, best_score,
            )
        else:
            point_assignments.append(-1)
            log.info(
                "[attach_images] point %d: no bag-of-words match (best_score=%d)",
                pi + 1, best_score,
            )

    # Второй проход: если пункт не получил картинку, но упоминает имя файла
    # группы в своём тексте — привязываем.  Минимум 1 совпадающий токен имени
    # файла (достаточно специфично: «CTO.CNt-008.1.25»).
    #
    # Распределяем глобально по убыванию score, чтобы пункт с самым сильным
    # совпадением (например, score=6 «СТО.CNt-008.1.25» в пункте 2) забрал
    # «свою» группу раньше пункта со score=1 (случайное совпадение токена).
    candidate_pairs: list[tuple[int, int, int]] = []  # (score, point_idx, group_idx)
    for i, point in enumerate(points):
        if point_assignments[i] >= 0:
            continue
        point_lower = point.lower()
        for idx, (_preceding_bag, _urls, fn) in enumerate(groups):
            if idx in used_groups or not fn:
                continue
            tokens = filename_tokens(fn)
            score = sum(1 for t in tokens if t in point_lower and len(t) >= 5)
            if score >= 1:
                candidate_pairs.append((score, i, idx))
    # Сортируем по убыванию score, при равенстве — по порядку появления пункта
    candidate_pairs.sort(key=lambda t: (-t[0], t[1]))
    for score, point_i, group_idx in candidate_pairs:
        if point_assignments[point_i] >= 0:
            continue
        if group_idx in used_groups:
            continue
        used_groups.add(group_idx)
        point_assignments[point_i] = group_idx
        log.info(
            "[attach_images] point %d → group %d (filename match score=%d, fn=%s)",
            point_i + 1, group_idx, score, groups[group_idx][2],
        )
    for i in range(len(points)):
        if point_assignments[i] < 0:
            log.info("[attach_images] point %d: no filename match", i + 1)

    # Финальная сборка с дедупликацией URL: одна и та же картинка может
    # оказаться в нескольких chunks (из-за chunk_overlap), из них
    # построятся разные groups с одинаковым URL, и разные пункты
    # ответа могут получить дубликат. Не выводим URL, который уже был.
    emitted_urls: set[str] = set()
    for point, group_idx in zip(points, point_assignments):
        if group_idx >= 0:
            urls = [u for u in groups[group_idx][1] if u not in emitted_urls]
            emitted_urls.update(urls)
            if urls:
                all_urls = "\n".join(urls)
                result_parts.append(f"{point}\n{all_urls}")
            else:
                result_parts.append(point)
        else:
            result_parts.append(point)

    return "\n".join(result_parts)


# ── Извлечение изображений из документов ─────────────────────────────────────

def _docx_has_ole_objects(path: Path) -> bool:
    """
    Проверяет, есть ли в .docx встроенные OLE-объекты (Visio-схемы и т.п.).
    OLE хранится как word/embeddings/oleObjectN.bin внутри docx-архива.
    Если есть — обычные средства python-docx не могут вытащить иллюстрацию.
    """
    import zipfile
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            for name in z.namelist():
                if name.startswith("word/embeddings/"):
                    return True
    except Exception:
        return False
    return False


def _convert_docx_to_pdf_via_word(docx_path: Path, pdf_path: Path) -> bool:
    """
    Конвертирует docx в PDF через установленный Microsoft Word (COM-автоматизация).
    Возвращает True при успехе.  Требует pywin32 и Windows-сервер с Word.

    Работает синхронно — вызывать через executor.  Word процесс закрывается
    в finally, даже при ошибке.
    """
    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except ImportError:
        log.warning(
            "[docx→pdf] pywin32 not installed — OLE objects in docx will be skipped. "
            "Install: pip install pywin32"
        )
        return False

    word = None
    document = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone
        document = word.Documents.Open(
            str(docx_path.resolve()),
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
        )
        # 17 = wdFormatPDF
        document.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        log.info("[docx→pdf] converted %s → %s via Word COM",
                 docx_path.name, pdf_path.name)
        return True
    except Exception as e:
        log.warning("[docx→pdf] Word COM conversion failed for %s: %s",
                    docx_path.name, e)
        return False
    finally:
        try:
            if document is not None:
                document.Close(SaveChanges=0)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _extract_emf_images_from_docx(
    docx_path: Path,
    images_dir: Path,
    start_index: int,
) -> list[str]:
    """
    Извлекает word/media/*.emf из docx-архива и конвертирует каждый файл
    в PNG через PowerShell + System.Drawing (нативно есть на Windows).

    EMF — это визуальные превью встроенных OLE-объектов (Visio схем,
    Excel-таблиц, Word-документов). Word гарантированно сохраняет такие
    превью при вставке OLE через "Insert Object", чтобы документ
    нормально открывался на ПК без оригинального приложения.

    Возвращает список имён созданных PNG (img_NNN.png) — каждое имя
    продолжает нумерацию с start_index + 1.

    Если PowerShell или System.Drawing недоступны (не-Windows),
    возвращает пустой список.
    """
    import zipfile
    import subprocess
    import tempfile

    created: list[str] = []
    try:
        with zipfile.ZipFile(str(docx_path), "r") as z:
            emf_names = sorted([
                n for n in z.namelist()
                if n.startswith("word/media/") and n.lower().endswith(".emf")
            ])
            if not emf_names:
                return []

            images_dir.mkdir(exist_ok=True)
            with tempfile.TemporaryDirectory() as tmpdir:
                tmpdir_path = Path(tmpdir)
                counter = start_index
                for emf_name in emf_names:
                    try:
                        emf_data = z.read(emf_name)
                    except Exception as e:
                        log.warning("[emf→png] failed to read %s: %s",
                                    emf_name, e)
                        continue
                    tmp_emf = tmpdir_path / Path(emf_name).name
                    try:
                        tmp_emf.write_bytes(emf_data)
                    except Exception as e:
                        log.warning("[emf→png] failed to write tmp %s: %s",
                                    emf_name, e)
                        continue

                    next_counter = counter + 1
                    png_name = f"img_{next_counter:03d}.png"
                    png_path = images_dir / png_name

                    ps_script = (
                        "Add-Type -AssemblyName System.Drawing; "
                        f"$img = [System.Drawing.Image]::FromFile('{tmp_emf}'); "
                        f"$img.Save('{png_path}', "
                        "[System.Drawing.Imaging.ImageFormat]::Png); "
                        "$img.Dispose()"
                    )
                    try:
                        result = subprocess.run(
                            [
                                "powershell",
                                "-NoProfile",
                                "-ExecutionPolicy", "Bypass",
                                "-Command", ps_script,
                            ],
                            capture_output=True,
                            timeout=30,
                            text=True,
                        )
                    except FileNotFoundError:
                        log.warning(
                            "[emf→png] powershell not found — EMF extraction "
                            "skipped for %s", docx_path.name
                        )
                        return created
                    except Exception as e:
                        log.warning("[emf→png] subprocess failed for %s: %s",
                                    emf_name, e)
                        continue

                    if result.returncode != 0 or not png_path.exists():
                        log.warning(
                            "[emf→png] conversion failed for %s "
                            "(rc=%d): %s",
                            emf_name, result.returncode,
                            (result.stderr or "")[:200].strip(),
                        )
                        continue

                    counter = next_counter
                    created.append(png_name)
                    log.info("[emf→png] %s → %s (%d bytes)",
                             emf_name, png_name, png_path.stat().st_size)
    except Exception as e:
        log.warning("[%s] EMF extraction failed: %s", docx_path.name, e)
    return created


def _extract_docx_via_pdf_render(path: Path) -> Optional[list]:
    """
    Стратегия для docx с OLE-объектами (Visio-схемы):
    1. Конвертирует docx в PDF через Word COM (Word уже учитывает страницы и
       рендерит OLE как векторную графику).
    2. Прогоняет полученный PDF через _extract_pdf_with_images — она уже
       умеет находить страницы с большой графикой и сохранять их как
       картинки (presentation mode).
    3. Возвращает результат с metadata.filename = имя исходного docx
       (чтобы download-ссылки и индексация работали с оригиналом).

    Возвращает None при неудаче — тогда вызывающий код использует
    обычный путь _extract_docx_with_images.
    """
    import tempfile
    from llama_index.core.schema import Document

    tmp_pdf = Path(tempfile.gettempdir()) / f"_docx2pdf_{path.stem}_{int(time.time())}.pdf"
    try:
        ok = _convert_docx_to_pdf_via_word(path, tmp_pdf)
        if not ok or not tmp_pdf.exists():
            return None

        # _extract_pdf_with_images создаст FILES_DIR/<stem>_images.
        # Чтобы папка картинок соответствовала имени docx, временно
        # переименуем PDF (точнее — передадим путь с именем docx).
        # Самый простой способ: подменим path.stem через симлинк/копию.
        # Здесь идём через копию имени: создаём вторую копию рядом.
        target_pdf = Path(tempfile.gettempdir()) / f"{path.stem}.pdf"
        try:
            if target_pdf.exists():
                target_pdf.unlink()
            tmp_pdf.replace(target_pdf)
        except Exception as e:
            log.warning("[docx→pdf] could not rename to %s: %s", target_pdf, e)
            target_pdf = tmp_pdf

        # ВАЖНО: _extract_pdf_with_images кладёт картинки в
        # FILES_DIR / f"{path.stem}_images".  Мы хотим, чтобы эта папка
        # совпала с именем исходного docx → даём ей путь с нужным stem.
        try:
            # docx→pdf от Word всегда даёт текстовый слой, но если в docx
            # один Visio на странице, текстового слоя в этой странице может
            # не быть — это нормально, не отвергаем такой PDF.
            pdf_docs = _extract_pdf_with_images(target_pdf, require_text_layer=False)
        finally:
            try:
                target_pdf.unlink(missing_ok=True)
            except Exception:
                pass

        # Извлекаем EMF-превью встроенных OLE-объектов (Visio схемы и
        # т.п.) напрямую из docx. Они дают точное визуальное
        # представление без угадывания по drawings.
        images_dir = FILES_DIR / f"{path.stem}_images"
        existing_imgs = []
        if images_dir.is_dir():
            existing_imgs = sorted(images_dir.glob("img_*.png"))
        start_idx = len(existing_imgs)
        emf_pngs = _extract_emf_images_from_docx(path, images_dir, start_idx)

        # Подменяем metadata.filename на имя исходного docx
        result = []
        for i, d in enumerate(pdf_docs):
            new_meta = dict(getattr(d, "metadata", {}) or {})
            new_meta["filename"] = path.name
            text = getattr(d, "text", "")
            # Маркеры EMF-картинок добавляем в конец текста первого
            # документа — attach_images найдёт их через extra image-chunks.
            # Каждая картинка проходит через vision-LLM (если включено),
            # и её описание встраивается рядом с маркером.
            if i == 0 and emf_pngs:
                markers = []
                for j, png_name in enumerate(emf_pngs):
                    idx = start_idx + j + 1
                    png_path = images_dir / png_name
                    markers.append(_build_image_marker(idx, png_name, png_path))
                text = text + "\n\n" + "\n".join(markers)
            result.append(Document(text=text, metadata=new_meta))
        log.info(
            "[%s] extracted via Word→PDF render: %d docs, +%d EMF images",
            path.name, len(result), len(emf_pngs),
        )
        return result
    except Exception as e:
        log.warning("[%s] Word→PDF pipeline failed: %s", path.name, e)
        return None
    finally:
        try:
            tmp_pdf.unlink(missing_ok=True)
        except Exception:
            pass


def _extract_docx_with_images(path: Path) -> list:
    """
    Извлекает текст из .docx с сохранением изображений.
    Изображения сохраняются в FILES_DIR/<stem>_images/img_001.png и т.д.
    В текст вставляются маркеры [Рисунок N: img_NNN.ext] на месте картинок.

    Если в docx есть OLE-объекты (например, Visio блок-схемы) — сначала
    пытаемся отрендерить через Word COM → PDF, потому что python-docx
    OLE-содержимое не видит.  При неудаче — обычный путь.
    """
    if _docx_has_ole_objects(path):
        log.info("[%s] docx contains OLE objects — trying Word→PDF render",
                 path.name)
        result = _extract_docx_via_pdf_render(path)
        if result is not None:
            return result
        log.info("[%s] falling back to python-docx (OLE objects will be lost)",
                 path.name)

    from docx import Document as DocxDocument
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from llama_index.core.schema import Document

    doc = DocxDocument(str(path))
    stem = path.stem
    images_dir = FILES_DIR / f"{stem}_images"
    images_dir.mkdir(exist_ok=True)

    # Собираем все inline-изображения с привязкой к параграфам
    img_counter = 0
    para_images: dict[int, list[str]] = {}  # para_index -> list of markers

    for para_idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawing_elements = run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            ) + run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
            )
            # Inline images через drawing
            blips = run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in blips:
                embed_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if not embed_id:
                    continue
                try:
                    rel = doc.part.rels[embed_id]
                except KeyError:
                    continue
                img_counter += 1
                img_data = rel.target_part.blob
                content_type = rel.target_part.content_type
                ext_map = {
                    "image/png": ".png",
                    "image/jpeg": ".jpg",
                    "image/gif": ".gif",
                    "image/bmp": ".bmp",
                    "image/tiff": ".tiff",
                    "image/x-wmf": ".wmf",
                    "image/x-emf": ".emf",
                }
                ext = ext_map.get(content_type, ".png")
                img_name = f"img_{img_counter:03d}{ext}"
                img_path = images_dir / img_name
                img_path.write_bytes(img_data)
                marker = _build_image_marker(img_counter, img_name, img_path)
                para_images.setdefault(para_idx, []).append(marker)
                log.info("[%s] Extracted image: %s", path.name, img_name)

    # Собираем текст с маркерами
    parts = []
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            parts.append(text)
        if para_idx in para_images:
            for marker in para_images[para_idx]:
                parts.append(marker)

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ValueError("No text extracted from DOCX file")

    # Удаляем папку, если картинок не нашлось
    if img_counter == 0:
        images_dir.rmdir()

    log.info("[%s] Extracted %d images", path.name, img_counter)
    return [Document(text=full_text, metadata={"filename": path.name})]


def _extract_pdf_with_images(path: Path, *, require_text_layer: bool = True) -> list:
    """
    Извлекает текст из .pdf с сохранением изображений через PyMuPDF (fitz).
    Изображения сохраняются в FILES_DIR/<stem>_images/img_001.ext и т.д.
    В текст вставляются маркеры [Рисунок N: img_NNN.ext] после текста каждой страницы.

    PyMuPDF корректно достаёт картинки любых типов (включая прозрачные, с масками,
    form XObjects), в отличие от pypdf — он молча пропускает сложные случаи.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(
            "PDF support requires PyMuPDF: pip install pymupdf"
        ) from e
    from llama_index.core.schema import Document

    doc = fitz.open(str(path))
    stem = path.stem
    images_dir = FILES_DIR / f"{stem}_images"
    images_dir.mkdir(exist_ok=True)

    img_counter = 0
    parts: list[str] = []
    seen_xrefs: dict[int, str] = {}  # один и тот же xref на разных страницах → один маркер

    # Pass 1: собираем xref-ы, которые используются как soft masks для других
    # изображений.  Soft mask — это служебная картинка, описывающая прозрачность
    # основной; вытащенная отдельно, выглядит как чёрный прямоугольник и не
    # имеет самостоятельного смысла.
    smask_xrefs: set[int] = set()
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            try:
                for img_info in page.get_images(full=True):
                    smask_xref = img_info[1] if len(img_info) > 1 else 0
                    if smask_xref:
                        smask_xrefs.add(int(smask_xref))
            except Exception:
                continue
    except Exception:
        pass
    if smask_xrefs:
        log.info("[%s] detected %d soft masks, will skip them", path.name, len(smask_xrefs))

    try:
        for page_num in range(len(doc)):
            page = doc[page_num]

            # Текст страницы
            page_text = page.get_text().strip()
            if page_text:
                parts.append(page_text)

            # Картинки страницы.  Вместо вытаскивания сырых байт через
            # extract_image() рендерим участок страницы по bbox каждой
            # картинки: PyMuPDF сам применяет soft masks, прозрачность,
            # gradients — и отдаёт ровно то, что видно в PDF-вьюере.
            # Для презентаций это единственный надёжный путь.
            try:
                image_info_list = page.get_image_info(xrefs=True)
            except Exception as e:
                log.warning("[%s] page %d get_image_info failed: %s",
                            path.name, page_num + 1, e)
                image_info_list = []

            # Шаг 1: собираем кандидатов с базовыми фильтрами
            #   - живой xref
            #   - не soft mask
            #   - не «иконка» (минимум 50 PDF-points по обеим сторонам)
            #   - не «крупная декорация» (>30% страницы без текста сверху)
            page_area = page.rect.width * page.rect.height
            candidates: list[tuple[int, tuple]] = []
            for info in image_info_list:
                xref = info.get("xref", 0)
                if not isinstance(xref, int) or xref <= 0:
                    continue
                if xref in smask_xrefs:
                    continue
                bbox = info.get("bbox")
                if not bbox or len(bbox) < 4:
                    continue
                try:
                    rect = fitz.Rect(bbox)
                    if rect.is_empty or rect.is_infinite:
                        continue
                    # Минимальный размер: ниже 150pt по одной из сторон —
                    # это всегда логотип / иконка / штамп, не контент.
                    # (Скриншоты UI и блок-схемы крупнее.)
                    if rect.width < 150 or rect.height < 150:
                        log.info(
                            "[%s] page %d skip small image xref=%d (%.0fx%.0f pt)",
                            path.name, page_num + 1, xref,
                            rect.width, rect.height,
                        )
                        continue
                    # Картинка в зоне header/footer (верхние 10% или нижние 10%
                    # страницы) — это колонтитул с логотипом, отбрасываем.
                    page_h = page.rect.height
                    if page_h > 0:
                        top_zone = 0.10 * page_h
                        bottom_zone = 0.90 * page_h
                        if rect.y1 <= top_zone:
                            log.info(
                                "[%s] page %d skip header image xref=%d",
                                path.name, page_num + 1, xref,
                            )
                            continue
                        if rect.y0 >= bottom_zone:
                            log.info(
                                "[%s] page %d skip footer image xref=%d",
                                path.name, page_num + 1, xref,
                            )
                            continue
                    # Крупная картинка без текста сверху — декорация (титульная
                    # иллюстрация, фон).  Картинки с текстом-наложением
                    # (кнопки, плашки, иконки с подписями) проходят фильтр.
                    img_area = rect.width * rect.height
                    if page_area > 0 and img_area / page_area > 0.30:
                        try:
                            text_over = page.get_text(clip=rect).strip()
                        except Exception:
                            text_over = ""
                        if not text_over:
                            log.info(
                                "[%s] page %d skip decorative xref=%d (%.0f%% area, no text)",
                                path.name, page_num + 1, xref,
                                100.0 * img_area / page_area,
                            )
                            continue
                except Exception:
                    continue
                candidates.append((xref, tuple(bbox)))

            page_markers: list[str] = []
            zoom = fitz.Matrix(2, 2)  # 2x для читаемости текста

            # Дополнительно: блок-схемы (Visio из docx, экспортированный через
            # Word) часто становятся не растровой картинкой, а большим набором
            # векторных операций (drawings).  Если такая графика занимает
            # заметную часть страницы — это блок-схема, рендерим страницу
            # целиком.
            # Автодетектор block-схем по drawings отключён: на типовых
            # Visio-схемах (белые блоки, чёрные контуры) ни один признак
            # не отличает их от пустых форм с таблицами. Visio-схемы
            # извлекаются напрямую из docx как EMF в _extract_docx_via_pdf_render.
            has_vector_graphic = False

            # Шаг 2: выбор стратегии — слайд или документ?
            # Слайд презентации:
            #   - либо много визуальных элементов (≥4),
            #   - либо мало текста (<500 символов) + хотя бы одна крупная
            #     картинка (>15% площади) — это «титульный» слайд с большой
            #     декоративной иллюстрацией (как стр.2 FAQ с большой монетой)
            #   - либо большая векторная графика (блок-схема из Visio)
            # Для слайдов рендерим всю страницу как одну картинку.
            # Для документных страниц (инструкция с парой скриншотов) —
            # каждую картинку отдельно.
            has_large_image = any(
                (b[2] - b[0]) * (b[3] - b[1]) > 0.15 * page_area
                for _, b in candidates
            ) if page_area > 0 else False
            page_text_len = len(page_text)
            is_presentation = (
                len(candidates) >= 4
                or (page_text_len < 500 and has_large_image)
                or has_vector_graphic
            )
            if is_presentation:
                try:
                    page_rect = page.rect
                    pix = page.get_pixmap(matrix=zoom, clip=page_rect, alpha=False)
                    img_bytes = pix.tobytes("png")
                    pix = None
                    img_counter += 1
                    img_name = f"img_{img_counter:03d}.png"
                    img_path = images_dir / img_name
                    img_path.write_bytes(img_bytes)
                    marker = _build_image_marker(img_counter, img_name, img_path)
                    page_markers.append(marker)
                    log.info(
                        "[%s] page %d: presentation mode (cands=%d, text=%d, big_img=%s, vector=%s) → 1 full-page render: %s",
                        path.name, page_num + 1, len(candidates),
                        page_text_len, has_large_image, has_vector_graphic, img_name,
                    )
                except Exception as e:
                    log.warning("[%s] page %d full-page render failed: %s",
                                path.name, page_num + 1, e)
            else:
                # Документная страница: каждая картинка отдельно
                for xref, bbox in candidates:
                    if xref in seen_xrefs:
                        page_markers.append(seen_xrefs[xref])
                        continue
                    try:
                        clip_rect = fitz.Rect(bbox)
                        pix = page.get_pixmap(matrix=zoom, clip=clip_rect, alpha=False)
                        img_bytes = pix.tobytes("png")
                        pix = None
                    except Exception as e:
                        log.warning("[%s] page %d render at xref=%d failed: %s",
                                    path.name, page_num + 1, xref, e)
                        continue

                    img_counter += 1
                    img_name = f"img_{img_counter:03d}.png"
                    img_path = images_dir / img_name
                    img_path.write_bytes(img_bytes)
                    marker = _build_image_marker(img_counter, img_name, img_path)
                    seen_xrefs[xref] = marker
                    page_markers.append(marker)
                    log.info("[%s] Rendered image from page %d: %s (xref=%d, bbox=%s)",
                             path.name, page_num + 1, img_name, xref,
                             tuple(round(v, 1) for v in bbox))

            for marker in page_markers:
                parts.append(marker)
        page_count = len(doc)
    finally:
        doc.close()

    full_text = "\n".join(parts).strip()

    # Считаем «полезный» текст — без маркеров [Рисунок N: img_NNN.ext].
    # Если его меньше PDF_MIN_TEXT_CHARS — это отсканированный PDF без
    # текстового слоя.  Распознавание сканов не поддерживается, отдаём
    # понятную ошибку.
    text_without_markers = re.sub(
        r"\[Рисунок\s+\d+:\s*[^\]]+\]", "", full_text
    ).strip()
    if require_text_layer and len(text_without_markers) < PDF_MIN_TEXT_CHARS:
        # подчищаем уже сохранённые картинки — файл всё равно отвергнут
        if images_dir.is_dir():
            try:
                shutil.rmtree(images_dir, ignore_errors=True)
            except Exception:
                pass
        raise ValueError(
            f"PDF выглядит как отсканированный — текстового слоя нет "
            f"(извлечено {len(text_without_markers)} симв. из {page_count} стр., "
            f"порог {PDF_MIN_TEXT_CHARS}). "
            "Распознавание сканов не поддерживается. "
            "Сохраните файл с текстовым слоем (например, через «Печать в PDF» "
            "из Word) и попробуйте снова."
        )

    if img_counter == 0:
        try:
            images_dir.rmdir()
        except OSError:
            pass

    log.info(
        "[%s] Extracted %d images from PDF (text chars: %d)",
        path.name, img_counter, len(text_without_markers),
    )
    return [Document(text=full_text, metadata={"filename": path.name})]


# ── Загрузка документов ───────────────────────────────────────────────────────

def _load_documents(path: Path):
    try:
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return _extract_docx_with_images(path)

        if suffix == ".pdf":
            return _extract_pdf_with_images(path)

        if suffix == ".msg":
            try:
                import extract_msg
                from llama_index.core.schema import Document
            except ImportError as e:
                raise RuntimeError(
                    "MSG support requires 'extract-msg' package"
                ) from e

            message = extract_msg.Message(str(path))
            parts = []
            subject = (message.subject or "").strip()
            sender = (message.sender or "").strip()
            date = (message.date or "").strip()
            body = (message.body or "").strip()
            if subject:
                parts.append(f"Subject: {subject}")
            if sender:
                parts.append(f"From: {sender}")
            if date:
                parts.append(f"Date: {date}")
            if body:
                parts.append("")
                parts.append(body)
            text = "\n".join(parts).strip()
            if not text:
                raise ValueError("No text extracted from MSG file")
            return [Document(text=text, metadata={"filename": path.name})]

        from llama_index.core import SimpleDirectoryReader

        return SimpleDirectoryReader(input_files=[str(path)]).load_data()
    finally:
        # Сохраняем накопленные описания картинок на диск, чтобы при
        # реиндексации не описывать те же самые повторно.
        _save_caption_cache()


def _detect_chunk_profile(texts: list[str]) -> tuple[str, int, int]:
    """
    Простейшая эвристика для mixed corpus:
    - cheatsheet: много коротких строк, списков и коротких заголовков
    - manual: более длинные связные абзацы
    """
    merged = "\n".join(t for t in texts if t).strip()
    if not merged:
        return ("manual", CHUNK_SIZE, CHUNK_OVERLAP)

    lines = [line.strip() for line in merged.splitlines() if line.strip()]
    if not lines:
        return ("manual", CHUNK_SIZE, CHUNK_OVERLAP)

    short_lines = sum(1 for line in lines if len(line) <= 80)
    numbered_lines = sum(1 for line in lines if re.match(r"^\d+[\.\)]\s+", line))
    heading_like = sum(
        1
        for line in lines
        if len(line) <= 120 and not line.endswith(".") and len(line.split()) <= 12
    )

    total_lines = len(lines)
    short_ratio = short_lines / total_lines
    numbered_ratio = numbered_lines / total_lines
    heading_ratio = heading_like / total_lines

    if short_ratio > 0.45 or numbered_ratio > 0.15 or heading_ratio > 0.20:
        return ("cheatsheet", 350, 50)
    return ("manual", 800, 120)


# ── Загрузка статей из корпоративной Базы знаний ─────────────────────────────

_IMG_CONTENT_TYPE_TO_EXT = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    "image/gif": "gif",
    "image/webp": "webp",
    "image/bmp": "bmp",
    "image/svg+xml": "svg",
}


def _ext_from_image_response(url: str, content_type: str) -> str:
    """Возвращает расширение картинки по content-type, с fallback на URL."""
    ct = (content_type or "").split(";", 1)[0].strip().lower()
    if ct in _IMG_CONTENT_TYPE_TO_EXT:
        return _IMG_CONTENT_TYPE_TO_EXT[ct]
    suffix = Path(url.split("?", 1)[0]).suffix.lstrip(".").lower()
    if suffix in {"png", "jpg", "jpeg", "gif", "webp", "bmp", "svg"}:
        return "jpg" if suffix == "jpeg" else suffix
    return "png"


async def _fetch_kb_article(url: str) -> tuple[str, str]:
    """
    Скачивает статью с портала Базы знаний (Yii + Apache Basic Auth).
    Возвращает (title, markdown_text).

    Авторизация через HTTP Basic — учётка прописана в KB_PORTAL_USER/PASSWORD.
    Apache на портале анонсирует и Negotiate (Kerberos), и Basic — мы идём
    по Basic-пути, потому что Negotiate требует ticket'ов AD на стороне Linux.

    Картинки из <img> качаются с теми же кредами и сохраняются в папку
    storage/files/<stem>_images/ с именами img_001.png и т.д.  В markdown
    на месте <img> вставляется маркер [Рисунок N: img_NNN.ext] — тот же
    формат, что используется для PDF/DOCX, чтобы существующая логика
    привязки картинок к ответам отработала автоматически.
    """
    if not KB_PORTAL_USER or not KB_PORTAL_PASSWORD:
        raise RuntimeError(
            "KB_PORTAL_USER and KB_PORTAL_PASSWORD must be configured in config.json"
        )
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise RuntimeError(
            "KB article support requires beautifulsoup4: pip install beautifulsoup4"
        ) from e
    from urllib.parse import urljoin

    auth = (KB_PORTAL_USER, KB_PORTAL_PASSWORD)
    async with httpx.AsyncClient(
        auth=auth,
        timeout=60,
        follow_redirects=True,
        verify=KB_PORTAL_VERIFY_SSL,
    ) as client:
        response = await client.get(url)
        if response.status_code == 401:
            raise ValueError(
                "Portal returned 401: wrong KB_PORTAL_USER/KB_PORTAL_PASSWORD"
            )
        response.raise_for_status()
        html = response.text

        soup = BeautifulSoup(html, "html.parser")

        # Title — внутри #pg-head лежит и кнопка «← Назад» (тег <a>),
        # и собственно заголовок статьи.  Убираем <a> ссылки и навигацию
        # перед извлечением текста — иначе title начинается со слова «Назад».
        title = ""
        title_el = soup.select_one("#pg-head")
        if title_el:
            # Копию делать не нужно: мы потом всё равно удаляем #pg-head
            # из body_el (см. блок чистки ниже), так что портим эту копию
            for nav in title_el.find_all(["a", "button", "nav"]):
                nav.decompose()
            title = title_el.get_text(separator=" ", strip=True)
        if not title:
            # Фоллбэк: первый h1/h2 в основном контенте, потом <title>
            for sel in ("#kb-b-container h1", "#kb-b-container h2", "h1"):
                el = soup.select_one(sel)
                if el and el.get_text(strip=True):
                    title = el.get_text(separator=" ", strip=True)
                    break
        if not title:
            title_tag = soup.find("title")
            title = title_tag.get_text(strip=True) if title_tag else "kb_article"
        # На случай если «Назад» где-то всё ещё в начале — режем
        title = re.sub(r"^(←\s*)?Назад\s+", "", title).strip() or "kb_article"

        # Body — основной контент в #kb-b-container
        body_el = soup.select_one("#kb-b-container")
        if not body_el:
            body_el = soup.select_one("#kb-container") or soup.select_one("main")
        if not body_el:
            raise ValueError(
                "Could not locate article body (#kb-b-container) in response HTML"
            )

        # Чистим тело: скрипты, стили, навигация, banner — это не контент
        for tag in body_el.select(
            "script, style, nav, .breadcrumb, #banner-cont, #pg-head"
        ):
            tag.decompose()

        # Скачиваем картинки и заменяем <img> на маркеры.
        # Папка-имя должна совпадать с тем, что вычислит upload_from_url
        # (она же используется эндпоинтом /documents/{filename}/images/...).
        # Важно: имя должно совпадать со stem'ом итогового файла,
        # включая префикс kb_<pid>_ если pid доступен.
        pid_for_dir = _extract_pid_from_url(url)
        if pid_for_dir is not None:
            stem = Path(_kb_filename_for_pid(pid_for_dir, title)).stem
        else:
            stem = _safe_title_to_filename(title)
        images_dir = FILES_DIR / f"{stem}_images"
        img_counter = 0

        for img in list(body_el.find_all("img")):
            src = (img.get("src") or "").strip()
            if not src or src.startswith("data:"):
                img.decompose()
                continue
            img_url = urljoin(url, src)
            if not img_url.startswith(("http://", "https://")):
                img.decompose()
                continue
            try:
                img_resp = await client.get(img_url)
                img_resp.raise_for_status()
            except Exception as e:
                log.warning("[KB] could not fetch image %s: %s", img_url, e)
                img.decompose()
                continue
            if img_counter == 0:
                images_dir.mkdir(exist_ok=True)
            img_counter += 1
            ext = _ext_from_image_response(
                img_url, img_resp.headers.get("content-type", "")
            )
            img_name = f"img_{img_counter:03d}.{ext}"
            img_path = images_dir / img_name
            img_path.write_bytes(img_resp.content)
            marker = _build_image_marker(img_counter, img_name, img_path)
            img.replace_with(marker)
            log.info("[KB] saved image %s ← %s", img_name, img_url)

    # Извлекаем текст с сохранением переносов параграфов.
    # img'и уже заменены на текстовые маркеры → попадают в текст естественно.
    text = body_el.get_text(separator="\n", strip=True)
    # Сжимаем подряд идущие пустые строки в одну
    text = re.sub(r"\n[ \t]*\n+", "\n\n", text).strip()

    if not text:
        raise ValueError("Article body is empty after extraction")

    log.info("[KB] extracted %d chars, %d images from %s", len(text), img_counter, url)

    # Собираем итоговый markdown: заголовок + тело
    markdown = f"# {title}\n\n{text}\n"
    return title, markdown


def _safe_title_to_filename(title: str) -> str:
    """Превращает заголовок статьи в безопасное имя файла."""
    cleaned = re.sub(r"[<>:\"/\\|?*\n\r\t]+", " ", title).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    if not cleaned:
        cleaned = f"kb_article_{int(time.time())}"
    return cleaned[:200]


def _extract_pid_from_url(url: str) -> Optional[int]:
    """Достаёт pid из URL вида .../kb/single?pid=189 — стабильный ID статьи."""
    m = re.search(r"[?&]pid=(\d+)", url)
    return int(m.group(1)) if m else None


def _kb_filename_for_pid(pid: int, title: str) -> str:
    """
    Имя файла для статьи Базы знаний.  Префикс kb_<pid>_ — стабильный ключ,
    не зависит от заголовка.  При переименовании статьи на портале pid тот же,
    поэтому старый файл легко найти по glob и заменить.
    """
    safe_title = _safe_title_to_filename(title)
    return f"kb_{pid:04d}_{safe_title}.md"


async def _remove_existing_kb_article(pid: int) -> int:
    """
    Удаляет любой существующий файл (и связанные с ним записи индекса
    и папку картинок) для статьи с этим pid.  Возвращает число удалённых файлов.
    """
    removed = 0
    for old_file in FILES_DIR.glob(f"kb_{pid:04d}_*.md"):
        log.info("[KB] removing previous version of pid=%d: %s", pid, old_file.name)
        try:
            await _delete_document_records(old_file.name)
        except Exception as e:
            log.warning("[KB] could not remove index records for %s: %s", old_file.name, e)
        try:
            old_file.unlink(missing_ok=True)
        except Exception as e:
            log.warning("[KB] could not unlink %s: %s", old_file, e)
        old_images = FILES_DIR / f"{old_file.stem}_images"
        if old_images.is_dir():
            try:
                shutil.rmtree(old_images, ignore_errors=True)
            except Exception as e:
                log.warning("[KB] could not rmtree %s: %s", old_images, e)
        _indexing.pop(old_file.name, None)
        removed += 1
    return removed


# ── Индексирование ────────────────────────────────────────────────────────────

async def _run_indexing(filename: str, dest: Path, category: Optional[str] = None):
    lock = asyncio.Lock()
    _indexing_locks[filename] = lock
    category = _normalize_category(category)

    async with lock:
        _indexing[filename] = {
            "status": "indexing",
            "chunks_done": 0,
            "chunks_total": 0,
            "error": None,
            "started_at": datetime.now(timezone.utc).isoformat(),
            "finished_at": None,
            "category": category,
        }
        try:
            from llama_index.core.node_parser import SentenceSplitter

            loop = asyncio.get_running_loop()
            docs = await loop.run_in_executor(None, lambda: _load_documents(dest))
            if not docs:
                raise ValueError("No text extracted from file")

            doc_texts = [getattr(doc, "text", "").strip() for doc in docs if getattr(doc, "text", "").strip()]
            profile, effective_chunk_size, effective_chunk_overlap = _detect_chunk_profile(doc_texts)
            log.info(
                "[%s] chunk profile: %s (%d/%d)",
                filename,
                profile,
                effective_chunk_size,
                effective_chunk_overlap,
            )

            splitter = SentenceSplitter(
                chunk_size=effective_chunk_size,
                chunk_overlap=effective_chunk_overlap,
            )
            nodes = await loop.run_in_executor(
                None,
                lambda: splitter.get_nodes_from_documents(docs),
            )

            texts = [node.text.strip() for node in nodes if node.text.strip()]
            total = len(texts)
            _indexing[filename]["chunks_total"] = total
            log.info("[%s] %d chunks (SentenceSplitter)", filename, total)

            done = 0
            new_records: list[dict] = []

            for batch_start in range(0, total, EMBED_BATCH_SIZE):
                batch_texts = texts[batch_start: batch_start + EMBED_BATCH_SIZE]
                embeddings = await _embed_batch(batch_texts)

                for offset, (text, emb) in enumerate(zip(batch_texts, embeddings)):
                    if emb is None:
                        log.warning("[%s] skipping chunk %d (embed 400)", filename, batch_start + offset)
                        continue
                    new_records.append(
                        {
                            "id": f"{filename}_{batch_start + offset}",
                            "filename": filename,
                            "chunk_index": batch_start + offset,
                            "text": text,
                            "embedding": _normalize_embedding(emb),
                            "category": category,
                        }
                    )
                    done += 1

                _indexing[filename]["chunks_done"] = done
                log.info("[%s] embedded %d/%d", filename, done, total)

            await _replace_document_records(filename, new_records)

            _indexing[filename].update(
                {
                    "status": "ready",
                    "chunks_done": done,
                    "finished_at": datetime.now(timezone.utc).isoformat(),
                }
            )
            log.info("[%s] done: %d chunks", filename, done)

        except Exception as e:
            log.error("[%s] failed: %s", filename, e)
            _indexing[filename].update(
                {
                    "status": "error",
                    "error": str(e),
                    "finished_at": datetime.now(timezone.utc).isoformat(),
                }
            )
        finally:
            _indexing_locks.pop(filename, None)


# ── FastAPI ───────────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    """
    При старте всегда перестраиваем FAISS из JSON — это дешевле, чем
    доверять файлу на диске, который мог рассинхронизироваться (в том числе
    при совпадении ntotal, но изменённом содержимом после сбоя).
    """
    global _faiss_index
    async with _records_lock:
        records = _load_records_sync()
        if records:
            log.info("Rebuilding FAISS index from JSON (%d records)...", len(records))
            _rebuild_faiss_index_sync(records)
        else:
            _faiss_index = None
            log.info("Index is empty, skipping FAISS build.")
    log.info("FAISS index ready: %d chunks", len(records))
    log.info("PDF policy: reject scans with text < %d chars", PDF_MIN_TEXT_CHARS)
    yield


app = FastAPI(title="RAG Ollama API v5.0", version="5.0.0", lifespan=lifespan)

# Статическая раздача файлов — прямые ссылки для 1С
app.mount("/files", StaticFiles(directory=str(FILES_DIR)), name="static_files")
app.mount("/prepared", StaticFiles(directory=str(PREPARED_DIR)), name="prepared_files")


class AskRequest(BaseModel):
    question: str
    top_k: Optional[int] = None
    category: Optional[str] = None


class AskResponse(BaseModel):
    answer: str
    answer_html: str
    sources: list[str]
    chunks_used: int
    download_urls: dict[str, str]
    image_urls: dict[str, str]
    raw_chunks: list[str]
    rewritten_query: Optional[str] = None


class UploadResponse(BaseModel):
    filename: str
    status: str
    category: Optional[str] = None


class PrepareResponse(BaseModel):
    source_filename: str
    prepared_filename: str
    status: str
    warnings: list[str]
    download_url: str
    text_url: str
    report_url: str


class DocumentInfo(BaseModel):
    filename: str
    chunks: int
    indexing_status: str
    download_url: str
    category: Optional[str] = None


class CategoryUpdateRequest(BaseModel):
    category: Optional[str] = None


class LoadFromUrlRequest(BaseModel):
    url: str
    category: Optional[str] = None


class IndexingStatus(BaseModel):
    filename: str
    status: str
    chunks_done: int
    chunks_total: int
    progress_pct: float
    error: Optional[str]
    started_at: Optional[str]
    finished_at: Optional[str]


class StatusResponse(BaseModel):
    total_chunks: int
    total_documents: int
    llm_model: str
    embed_model: str
    ollama_url: str


@app.get("/status", response_model=StatusResponse)
async def status():
    total = await _chunk_count()
    docs = len(await _document_chunk_counts())
    return StatusResponse(
        total_chunks=total,
        total_documents=docs,
        llm_model=OLLAMA_LLM_MODEL,
        embed_model=OLLAMA_EMBED_MODEL,
        ollama_url=OLLAMA_BASE_URL,
    )


@app.get("/documents", response_model=list[DocumentInfo])
async def list_documents(request: Request, category: Optional[str] = None):
    counts = await _document_chunk_counts()
    cats = await _document_categories()
    for fn in _indexing:
        counts.setdefault(fn, 0)
    base = _public_base_url(request)
    cat_filter = _normalize_category(category)
    result = []
    for fn, n in sorted(counts.items()):
        doc_cat = cats.get(fn) or _indexing.get(fn, {}).get("category")
        if cat_filter is not None and doc_cat != cat_filter:
            continue
        st = _indexing.get(fn, {}).get("status", "ready" if n > 0 else "unknown")
        result.append(
            DocumentInfo(
                filename=fn,
                chunks=n,
                indexing_status=st,
                download_url=f"{base}/files/{fn}",
                category=doc_cat,
            )
        )
    return result


@app.get("/categories", response_model=list[str])
async def list_categories():
    return await _list_categories()


@app.get("/logs")
async def get_logs(
    limit: int = 200,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
):
    """
    Возвращает последние записи лога /ask, новые сверху.

    - limit: 1..LOG_MAX_RETURN
    - date_from / date_to: ISO-строка ("2026-04-30" или "2026-04-30T12:00:00Z").
      Сравнение лексикографическое — для ISO-формата это эквивалентно сравнению дат.
    """
    if not isinstance(limit, int) or limit < 1 or limit > LOG_MAX_RETURN:
        raise HTTPException(400, f"limit must be 1..{LOG_MAX_RETURN}, got: {limit}")
    return await _read_ask_log(limit=limit, date_from=date_from, date_to=date_to)


@app.post("/documents", response_model=UploadResponse, status_code=202)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
):
    filename = _safe_filename(file.filename)
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED:
        raise HTTPException(
            400,
            f"Unsupported type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED))}",
        )
    if _indexing.get(filename, {}).get("status") == "indexing":
        raise HTTPException(409, f"'{filename}' is already being indexed")
    dest = FILES_DIR / filename
    async with aiofiles.open(dest, "wb") as f:
        while chunk := await file.read(1024 * 256):
            await f.write(chunk)
    norm_cat = _normalize_category(category)
    background_tasks.add_task(_run_indexing, filename, dest, norm_cat)
    return UploadResponse(filename=filename, status="indexing_started", category=norm_cat)


@app.post("/documents/from-url", response_model=UploadResponse, status_code=202)
async def upload_from_url(background_tasks: BackgroundTasks, body: LoadFromUrlRequest):
    """
    Скачивает статью с корпоративной Базы знаний по URL, сохраняет как
    markdown-файл и запускает индексацию.  Авторизация на портале — HTTP Basic
    с учёткой из config.json (KB_PORTAL_USER / KB_PORTAL_PASSWORD).
    """
    url = (body.url or "").strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(400, "url must start with http:// or https://")

    # ВАЖНО: сначала удаляем старую версию (вместе с её папкой картинок),
    # ТОЛЬКО ПОТОМ скачиваем новую — иначе _remove_existing_kb_article
    # снесёт свежескачанную папку kb_<pid>_..._images.
    pid = _extract_pid_from_url(url)
    if pid is not None:
        removed = await _remove_existing_kb_article(pid)
        if removed:
            log.info("[KB] replaced %d previous version(s) of pid=%d", removed, pid)

    try:
        title, markdown = await _fetch_kb_article(url)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except httpx.HTTPStatusError as e:
        raise HTTPException(502, f"Portal returned HTTP {e.response.status_code}")
    except httpx.RequestError as e:
        raise HTTPException(502, f"Could not reach portal: {e}")
    except RuntimeError as e:
        raise HTTPException(500, str(e))

    if pid is not None:
        filename = _safe_filename(_kb_filename_for_pid(pid, title))
    else:
        safe_title = _safe_title_to_filename(title)
        filename = _safe_filename(f"{safe_title}.md")

    if _indexing.get(filename, {}).get("status") == "indexing":
        raise HTTPException(409, f"'{filename}' is already being indexed")

    dest = FILES_DIR / filename
    async with aiofiles.open(dest, "w", encoding="utf-8") as f:
        await f.write(markdown)
    log.info("Saved KB article from %s → %s (%d chars)", url, filename, len(markdown))

    norm_cat = _normalize_category(body.category)
    background_tasks.add_task(_run_indexing, filename, dest, norm_cat)
    return UploadResponse(filename=filename, status="indexing_started", category=norm_cat)


async def _bulk_import_kb_articles(
    pid_min: int,
    pid_max: int,
    category: Optional[str],
) -> None:
    """
    Фоновая задача: перебирает pid от pid_min до pid_max включительно,
    скачивает каждую существующую статью и ставит на индексацию.
    Несуществующие pid'ы (404 / 401 / редиректы на /kb/pg) пропускаются.
    """
    log.info("[KB bulk] starting brute-force import pid=%d..%d", pid_min, pid_max)
    base = KB_PORTAL_BASE_URL.rstrip("/")
    started = 0
    skipped = 0
    failed = 0
    for pid in range(pid_min, pid_max + 1):
        url = f"{base}/kb/single?pid={pid}"
        # ВАЖНО: сначала удаляем старую версию (вместе с папкой картинок),
        # потом скачиваем новую — иначе _remove_existing_kb_article снесёт
        # свежескачанную папку kb_<pid>_..._images.
        try:
            await _remove_existing_kb_article(pid)
        except Exception as e:
            log.warning("[KB bulk] pid=%d remove-old error: %s", pid, e)
        try:
            title, markdown = await _fetch_kb_article(url)
        except ValueError as e:
            # пустое тело, нет #kb-b-container и т.п. — нет такого pid
            skipped += 1
        except httpx.HTTPStatusError as e:
            if e.response.status_code in (401, 403, 404):
                skipped += 1
            else:
                log.warning("[KB bulk] pid=%d HTTP %d", pid, e.response.status_code)
                failed += 1
        except httpx.RequestError as e:
            log.warning("[KB bulk] pid=%d network error: %s", pid, e)
            failed += 1
        except Exception as e:
            log.warning("[KB bulk] pid=%d unexpected: %s", pid, e)
            failed += 1
        else:
            try:
                # Стабильное имя kb_<pid>_<title>.md
                filename = _safe_filename(_kb_filename_for_pid(pid, title))
                if _indexing.get(filename, {}).get("status") == "indexing":
                    log.info("[KB bulk] pid=%d already indexing %s, skip", pid, filename)
                    skipped += 1
                else:
                    dest = FILES_DIR / filename
                    async with aiofiles.open(dest, "w", encoding="utf-8") as f:
                        await f.write(markdown)
                    # fire-and-forget — не блокируемся на индексации
                    asyncio.create_task(_run_indexing(filename, dest, category))
                    started += 1
                    log.info("[KB bulk] pid=%d saved → %s", pid, filename)
            except Exception as e:
                log.warning("[KB bulk] pid=%d save error: %s", pid, e)
                failed += 1

        # быть вежливыми к порталу
        await asyncio.sleep(0.3)

    log.info(
        "[KB bulk] done: pid=%d..%d, started=%d, skipped=%d, failed=%d",
        pid_min, pid_max, started, skipped, failed,
    )


class BulkImportRequest(BaseModel):
    pid_min: int = 1
    pid_max: int = 500
    category: Optional[str] = None


@app.post("/documents/from-portal-bulk", status_code=202)
async def upload_from_portal_bulk(
    background_tasks: BackgroundTasks, body: BulkImportRequest
):
    """
    Запускает фоновый импорт всех статей с корпоративной Базы знаний.
    Перебирает pid от pid_min до pid_max, скачивает существующие статьи
    и ставит на индексацию.  Возвращается сразу с 202; прогресс — в логе.
    """
    if body.pid_min < 1 or body.pid_max < body.pid_min or body.pid_max > 10_000:
        raise HTTPException(400, "pid_min/pid_max must satisfy 1 ≤ min ≤ max ≤ 10000")

    norm_cat = _normalize_category(body.category)
    background_tasks.add_task(
        _bulk_import_kb_articles, body.pid_min, body.pid_max, norm_cat,
    )
    return {
        "status": "bulk_import_started",
        "pid_min": body.pid_min,
        "pid_max": body.pid_max,
        "category": norm_cat,
    }


@app.post("/documents/prepare", response_model=PrepareResponse)
async def prepare_document(
    request: Request,
    file: UploadFile = File(...),
):
    filename = _safe_filename(file.filename)
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED:
        raise HTTPException(
            400,
            f"Unsupported type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED))}",
        )

    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S%fZ")
    raw_path = PREPARED_DIR / f"source_{ts}_{filename}"
    async with aiofiles.open(raw_path, "wb") as f:
        while chunk := await file.read(1024 * 256):
            await f.write(chunk)

    images_dir = FILES_DIR / f"{raw_path.stem}_images"
    if images_dir.is_dir():
        shutil.rmtree(images_dir, ignore_errors=True)

    try:
        loop = asyncio.get_running_loop()
        docs = await loop.run_in_executor(None, lambda: _load_documents(raw_path))
        source_text = "\n\n".join(
            getattr(doc, "text", "").strip()
            for doc in docs
            if getattr(doc, "text", "").strip()
        )
        if not source_text:
            raise ValueError("No text extracted from file")

        prepared_text = await _normalize_text_for_rag(source_text)
        warnings = _validate_prepared_text(prepared_text, filename, source_text=source_text)
    except Exception as e:
        log.error("[%s] prepare failed: %s", filename, e)
        raise HTTPException(500, f"Document prepare failed: {e}") from e
    finally:
        raw_path.unlink(missing_ok=True)

    prepared_name = _prepared_filename(filename)
    prepared_path = PREPARED_DIR / prepared_name
    images_dir = FILES_DIR / f"{raw_path.stem}_images"
    await asyncio.to_thread(_write_prepared_docx, prepared_text, prepared_path, images_dir)

    prepared_text_name = f"{Path(prepared_name).stem}.md"
    prepared_text_path = PREPARED_DIR / prepared_text_name
    prepared_text_path.write_text(prepared_text, encoding="utf-8")

    if images_dir.is_dir():
        shutil.rmtree(images_dir, ignore_errors=True)

    report_name = f"{Path(prepared_name).stem}_report.json"
    report_path = PREPARED_DIR / report_name
    report = {
        "source_filename": filename,
        "prepared_filename": prepared_name,
        "status": "ok" if not warnings else "warnings",
        "warnings": warnings,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "llm_model": OLLAMA_LLM_MODEL,
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    base = _public_base_url(request)
    return PrepareResponse(
        source_filename=filename,
        prepared_filename=prepared_name,
        status=report["status"],
        warnings=warnings,
        download_url=f"{base}/prepared/{prepared_name}",
        text_url=f"{base}/prepared/{prepared_text_name}",
        report_url=f"{base}/prepared/{report_name}",
    )


@app.get("/documents/{filename}/download")
async def download_document(filename: str):
    filename = _safe_filename(filename)
    path = FILES_DIR / filename
    if not path.exists():
        raise HTTPException(404, f"File '{filename}' not found")
    return FileResponse(path, filename=filename)


@app.get("/documents/{filename}/images/{image_name}")
async def download_image(filename: str, image_name: str):
    filename = _safe_filename(filename)
    image_name = Path(image_name).name
    if not image_name:
        raise HTTPException(400, "Invalid image name")
    stem = Path(filename).stem
    img_path = FILES_DIR / f"{stem}_images" / image_name
    if not img_path.exists():
        raise HTTPException(404, f"Image '{image_name}' not found for '{filename}'")
    return FileResponse(img_path, filename=image_name)


@app.get("/documents/{filename}/status", response_model=IndexingStatus)
async def document_status(filename: str):
    filename = _safe_filename(filename)
    info = _indexing.get(filename)
    if info is None:
        chunks = len(await _document_records(filename))
        if chunks:
            return IndexingStatus(
                filename=filename,
                status="ready",
                chunks_done=chunks,
                chunks_total=chunks,
                progress_pct=100.0,
                error=None,
                started_at=None,
                finished_at=None,
            )
        raise HTTPException(404, f"Document '{filename}' not found")
    total = info["chunks_total"] or 1
    return IndexingStatus(
        filename=filename,
        status=info["status"],
        chunks_done=info["chunks_done"],
        chunks_total=info["chunks_total"],
        progress_pct=round(info["chunks_done"] / total * 100, 1),
        error=info.get("error"),
        started_at=info.get("started_at"),
        finished_at=info.get("finished_at"),
    )


@app.post("/documents/{filename}/reindex")
async def reindex_document(
    filename: str,
    background_tasks: BackgroundTasks,
    category: Optional[str] = None,
):
    """Переиндексировать существующий документ без повторной загрузки.

    Если category не задана — сохраняем существующую категорию документа.
    Чтобы явно сбросить категорию, передайте пустую строку.
    """
    filename = _safe_filename(filename)
    dest = FILES_DIR / filename
    if not dest.exists():
        raise HTTPException(404, f"File '{filename}' not found")
    if _indexing.get(filename, {}).get("status") == "indexing":
        raise HTTPException(409, f"'{filename}' is already being indexed")
    if category is None:
        effective_cat = await _document_category(filename)
    else:
        effective_cat = _normalize_category(category)
    # Удаляем старые чанки и картинки.
    # ВАЖНО: для .md картинки НЕ трогаем — при реиндексации .md-файла
    # его текст только читается с диска, картинки заново не скачиваются
    # (в отличие от PDF/docx, которые пересобирают картинки из документа).
    # Стереть папку тут = получить битые ссылки на картинки в ответе.
    await _delete_document_records(filename)
    if dest.suffix.lower() != ".md":
        images_dir = FILES_DIR / f"{Path(filename).stem}_images"
        if images_dir.is_dir():
            shutil.rmtree(images_dir, ignore_errors=True)
    background_tasks.add_task(_run_indexing, filename, dest, effective_cat)
    return {"filename": filename, "status": "reindexing_started", "category": effective_cat}


@app.post("/documents/reindex-all")
async def reindex_all(background_tasks: BackgroundTasks):
    """Переиндексировать все документы (категории сохраняются)."""
    files = [f for f in FILES_DIR.iterdir() if f.is_file() and f.suffix.lower() in SUPPORTED]
    if not files:
        raise HTTPException(404, "No documents found")
    cats = await _document_categories()
    started = []
    skipped = []
    for dest in files:
        fn = dest.name
        if _indexing.get(fn, {}).get("status") == "indexing":
            skipped.append(fn)
            continue
        existing_cat = cats.get(fn)
        await _delete_document_records(fn)
        # Для .md не сносим картинки — они не пересоздаются при
        # реиндексации файла с диска (см. reindex_document).
        if dest.suffix.lower() != ".md":
            images_dir = FILES_DIR / f"{dest.stem}_images"
            if images_dir.is_dir():
                shutil.rmtree(images_dir, ignore_errors=True)
        background_tasks.add_task(_run_indexing, fn, dest, existing_cat)
        started.append(fn)
    return {"started": started, "skipped_already_indexing": skipped}


@app.put("/documents/{filename}/category")
async def update_document_category(filename: str, body: CategoryUpdateRequest):
    """Изменить категорию документа без переиндексации."""
    filename = _safe_filename(filename)
    if not await _document_records(filename):
        raise HTTPException(404, f"Document '{filename}' not found in index")
    norm_cat = _normalize_category(body.category)
    updated = await _update_document_category(filename, norm_cat)
    return {"filename": filename, "category": norm_cat, "chunks_updated": updated}


@app.delete("/documents")
async def delete_documents_by_category(category: Optional[str] = None):
    """Удалить все документы указанной категории."""
    norm_cat = _normalize_category(category)
    if norm_cat is None:
        raise HTTPException(400, "category is required")

    indexing_now = [
        fn for fn, info in _indexing.items()
        if info.get("status") == "indexing" and info.get("category") == norm_cat
    ]
    if indexing_now:
        raise HTTPException(
            409,
            f"Category '{norm_cat}' has documents currently being indexed: {', '.join(indexing_now)}",
        )

    removed_by_file = await _delete_category_records(norm_cat)
    if not removed_by_file:
        raise HTTPException(404, f"No documents found for category '{norm_cat}'")

    for filename in removed_by_file:
        (FILES_DIR / filename).unlink(missing_ok=True)
        images_dir = FILES_DIR / f"{Path(filename).stem}_images"
        if images_dir.is_dir():
            shutil.rmtree(images_dir, ignore_errors=True)
        _indexing.pop(filename, None)

    return {
        "category": norm_cat,
        "documents_removed": sorted(removed_by_file),
        "chunks_removed": sum(removed_by_file.values()),
    }


@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    filename = _safe_filename(filename)
    if _indexing.get(filename, {}).get("status") == "indexing":
        raise HTTPException(409, f"'{filename}' is currently being indexed")
    removed = await _delete_document_records(filename)
    if not removed:
        raise HTTPException(404, f"Document '{filename}' not found in index")
    (FILES_DIR / filename).unlink(missing_ok=True)
    # Удаляем папку с извлечёнными изображениями
    images_dir = FILES_DIR / f"{Path(filename).stem}_images"
    if images_dir.is_dir():
        shutil.rmtree(images_dir, ignore_errors=True)
    _indexing.pop(filename, None)
    return {"filename": filename, "chunks_removed": removed}


@app.post("/ask", response_model=AskResponse)
async def ask(req: AskRequest, request: Request):
    log_entry: dict = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "question": req.question,
        "category": req.category,
        "top_k": req.top_k,
        "best_similarity": None,
        "sources": [],
        "chunks_used": 0,
        "rewritten_query": None,
        "answer": None,
        "error": None,
        "duration_ms": None,
    }
    started_mono = time.monotonic()

    try:
        if await _chunk_count() == 0:
            raise HTTPException(503, "Index is empty. Upload documents first via POST /documents")

        raw_top_k = req.top_k
        if raw_top_k is not None and (not isinstance(raw_top_k, int) or raw_top_k < 1):
            raise HTTPException(400, f"top_k must be a positive integer, got: {raw_top_k}")
        top_k = max(1, raw_top_k or TOP_K)

        rewritten = await _rewrite_query(req.question)
        log.info("Query rewrite: %r → %r", req.question, rewritten)
        log_entry["rewritten_query"] = rewritten

        try:
            q_emb = await _embed(rewritten)
            if q_emb is None:
                raise ValueError("embed returned None")
        except Exception as e:
            raise HTTPException(502, f"Embedding API error: {e}")

        category = _normalize_category(req.category)
        results = await _search_records(rewritten, _normalize_embedding(q_emb), top_k, category=category)

        docs = results["documents"][0]
        metas = results["metadatas"][0]
        distances = results["distances"][0]
        filename_scores = results.get("filename_scores", [[]])[0]

        _not_found = "Информация в документах не найдена."
        _not_found_html = _answer_to_html(_not_found, {})

        if not docs:
            log_entry["answer"] = _not_found
            return AskResponse(
                answer=_not_found,
                answer_html=_not_found_html,
                sources=[],
                chunks_used=0,
                download_urls={},
                image_urls={},
                raw_chunks=[],
                rewritten_query=rewritten,
            )

        best_sim = 1 - distances[0]
        best_filename_score = filename_scores[0] if filename_scores else 0.0
        # distance ≥ 1.0 означает, что top-результат пришёл из BM25 (vector
        # search его не нашёл) — это валидный лексический хит, к нему
        # не применяем порог по vector similarity.
        is_bm25_only = distances[0] >= 1.0
        log.info(
            "Best similarity: %.3f (filename_score=%.3f, bm25_only=%s)",
            best_sim, best_filename_score, is_bm25_only,
        )
        log_entry["best_similarity"] = round(float(best_sim), 4)
        if best_sim < SIM_THRESHOLD and best_filename_score < 0.7 and not is_bm25_only:
            log_entry["answer"] = _not_found
            return AskResponse(
                answer=_not_found,
                answer_html=_not_found_html,
                sources=[],
                chunks_used=0,
                download_urls={},
                image_urls={},
                raw_chunks=[],
                rewritten_query=rewritten,
            )

        # Расширяем контекст соседними чанками (document-aware)
        records = await _get_records()
        all_docs, all_metas = _build_document_aware_context(
            results["ids"][0], records, budget=top_k * 2,
        )

        if not all_docs:
            all_docs = docs
            all_metas = metas

        # Дедуплицируем чанки по нормализованному тексту: один и тот же
        # фрагмент мог попасть и через vector, и через BM25, и через
        # соседние чанки document-aware. Иначе LLM выдаёт одну и ту же
        # фразу N раз в нумерованном списке.
        _seen_norm: set[str] = set()
        _dedup_docs: list[str] = []
        _dedup_metas: list[dict] = []
        for _d, _m in zip(all_docs, all_metas):
            _norm = re.sub(r"\s+", " ", _d).strip().lower()
            if not _norm or _norm in _seen_norm:
                continue
            _seen_norm.add(_norm)
            _dedup_docs.append(_d)
            _dedup_metas.append(_m)
        if len(_dedup_docs) < len(all_docs):
            log.info(
                "Deduplicated chunks: %d → %d", len(all_docs), len(_dedup_docs)
            )
        all_docs, all_metas = _dedup_docs, _dedup_metas

        context = "\n\n".join(
            f"[CHUNK {i} | {m['filename']}]\n{d}"
            for i, (d, m) in enumerate(zip(all_docs, all_metas))
        )
        relevant_fragments = _most_relevant_fragments(req.question, all_docs, all_metas)
        relevant_context = "\n".join(
            f"{i}. {fragment}"
            for i, fragment in enumerate(relevant_fragments, start=1)
        )
        # Широкий запрос (короткий, без конкретики «как/где/ошибка/...») — это,
        # как правило, название инструкции целиком. В таком случае выделять
        # один «целевой раздел» вредно: получим один пункт из середины.
        # Лучше отдать LLM полный контекст и разрешить объединять разделы.
        is_broad = _is_broad_query(req.question)
        if is_broad:
            target_context = ""
        else:
            target_context = _target_section_context(req.question, all_docs, all_metas)
        _all_sources = list(dict.fromkeys(m["filename"] for m in all_metas))

        if is_broad:
            system_content = (
                "Ты — помощник сотрудников ГК ОСНОВА. Отвечай на русском живо, "
                "по-человечески, как коллега объясняет коллеге. Не в стиле "
                "служебной записки.\n"
                "\n"
                "СТРУКТУРА:\n"
                "Первая строка — вводное предложение о сути темы (одно, максимум "
                "два). Сразу к делу: что это, зачем нужно, кому касается. "
                "Своими словами, не копия фразы из контекста.\n"
                "Пустая строка.\n"
                "Дальше нумерованный список 1. 2. 3. ... Каждый пункт — 2–3 полных "
                "предложения, живой связный текст. Не обрывок, не рубленая фраза, "
                "не заголовок.\n"
                "\n"
                "ПРИМЕР 1 (про валюту):\n"
                "«СНОВАКСы» — внутренняя валюта ГК ОСНОВА для поощрения полезных "
                "действий сотрудников. Ей можно расплачиваться на корпоративном "
                "маркетплейсе.\n"
                "\n"
                "1. Заработать «СНОВАКСы» может любой сотрудник, независимо от "
                "должности и юрлица. Для этого нужно активно участвовать в жизни "
                "компании — быть пунктуальным, публиковать идеи, участвовать в "
                "опросах.\n"
                "2. Начисления происходят автоматически: за каждое действие "
                "система переводит валюту на электронный кошелёк. Баланс всегда "
                "виден в личном кабинете на корпоративном портале.\n"
                "3. Потратить накопленное можно на маркетплейсе — там доступны "
                "товары, услуги и внутренние привилегии от партнёров.\n"
                "\n"
                "ПРИМЕР 2 (про процесс):\n"
                "Отпуск в ГК ОСНОВА оформляется через 1С:Документооборот и "
                "требует предварительного согласования руководителя.\n"
                "\n"
                "1. Сотрудник создаёт заявку на отпуск в 1С минимум за две недели "
                "до желаемой даты. Указывает период, тип отпуска и остаток дней.\n"
                "2. Заявка автоматически уходит на согласование руководителю "
                "подразделения. Статус можно отслеживать в карточке заявки.\n"
                "3. После согласования отпуск попадает в общий график, а сотруднику "
                "приходит уведомление на корпоративную почту.\n"
                "\n"
                "ФАКТЫ:\n"
                "Все факты — только из контекста ниже. Формулировки меняй "
                "свободно, факты не выдумывай: числа, сроки, роли, шаги — как в "
                "исходнике.\n"
                "\n"
                "Вопрос общий — объединяй материал из разных документов. Не "
                "смешивай факты про разные сущности.\n"
                "\n"
                "Если в контексте нет сути вопроса — ответь ровно одной фразой: "
                "Информация в документах не найдена.\n"
                "\n"
                "ЧТО ТОЧНО НЕ ДЕЛАТЬ:\n"
                "— Не начинай введение с «Тема этого материала...», «Данный "
                "документ...», «Здесь описывается...», «Давайте разберём», «Вот "
                "пошаговая инструкция», «Okay».\n"
                "— Не пиши пункты одним предложением-обрубком. Каждый пункт — "
                "минимум два-три полных предложения с пояснением.\n"
                "— Не переводи термины на английский.\n"
                "— Не используй LaTeX и $…$: стрелка — символ →.\n"
                "— Строки «Изображение: ...» — это описания картинок с ошибками "
                "распознавания. Слова, которые встречаются только там и больше "
                "нигде в контексте, в ответ не переноси."
            )
        else:
            system_content = (
                "Ты — помощник сотрудников ГК ОСНОВА. Отвечай на русском живо, "
                "по-человечески, как коллега объясняет коллеге. Не в стиле "
                "служебной записки.\n"
                "\n"
                "СТРУКТУРА:\n"
                "Первая строка — вводное предложение о сути темы (одно, максимум "
                "два). Сразу к делу: что это, зачем нужно, кому касается. "
                "Своими словами, не копия фразы из контекста.\n"
                "Пустая строка.\n"
                "Дальше нумерованный список 1. 2. 3. ... Каждый пункт — 2–3 полных "
                "предложения, живой связный текст. Не обрывок, не рубленая фраза, "
                "не заголовок.\n"
                "\n"
                "ПРИМЕР 1 (про процесс):\n"
                "Отпуск в ГК ОСНОВА оформляется через 1С:Документооборот и "
                "требует предварительного согласования руководителя.\n"
                "\n"
                "1. Сотрудник создаёт заявку на отпуск в 1С минимум за две недели "
                "до желаемой даты. Указывает период, тип отпуска и остаток дней.\n"
                "2. Заявка автоматически уходит на согласование руководителю "
                "подразделения. Статус можно отслеживать в карточке заявки.\n"
                "3. После согласования отпуск попадает в общий график, а сотруднику "
                "приходит уведомление на корпоративную почту.\n"
                "\n"
                "ПРИМЕР 2 (про справочник/термин):\n"
                "«СНОВАКСы» — внутренняя валюта ГК ОСНОВА для поощрения полезных "
                "действий сотрудников. Ей можно расплачиваться на корпоративном "
                "маркетплейсе.\n"
                "\n"
                "1. Заработать «СНОВАКСы» может любой сотрудник, независимо от "
                "должности и юрлица. Начисляются они за пунктуальность, участие в "
                "мероприятиях и публикацию идей.\n"
                "2. Баланс отражается в личном кабинете на корпоративном портале "
                "и обновляется автоматически после каждого действия.\n"
                "3. Потратить накопленное можно на маркетплейсе — там доступны "
                "товары, услуги и внутренние привилегии.\n"
                "\n"
                "ФАКТЫ:\n"
                "Все факты — только из контекста ниже. Формулировки меняй "
                "свободно, факты не выдумывай: числа, сроки, роли, шаги — как в "
                "исходнике.\n"
                "\n"
                "Вопрос конкретный — отвечай строго по разделу, который прямо "
                "соответствует вопросу. Не подмешивай соседние разделы, если они о "
                "другой теме. Исправления, примечания и подсказки — только из "
                "того же раздела.\n"
                "\n"
                "Если блок «Целевой раздел» заполнен — отвечай только по нему; "
                "полный контекст используй как справочник для проверки.\n"
                "\n"
                "Не переставляй шаги местами и не начинай с середины процесса.\n"
                "\n"
                "Если в контексте нет сути вопроса — ответь ровно одной фразой: "
                "Информация в документах не найдена.\n"
                "\n"
                "ЧТО ТОЧНО НЕ ДЕЛАТЬ:\n"
                "— Не начинай введение с «Тема этого материала...», «Данный "
                "документ...», «Здесь описывается...», «Давайте разберём», «Вот "
                "пошаговая инструкция», «Okay».\n"
                "— Не пиши пункты одним предложением-обрубком. Каждый пункт — "
                "минимум два-три полных предложения с пояснением.\n"
                "— Не переводи термины на английский.\n"
                "— Не используй LaTeX и $…$: стрелка — символ →.\n"
                "— Строки «Изображение: ...» — это описания картинок с ошибками "
                "распознавания. Слова, которые встречаются только там и больше "
                "нигде в контексте, в ответ не переноси."
            )
        log.info("Query mode: %s", "broad" if is_broad else "specific")

        try:
            answer = await _chat(
                [
                    {
                        "role": "system",
                        "content": system_content,
                    },
                    {
                        "role": "user",
                        "content": (
                            f"Наиболее релевантные фрагменты:\n\n"
                            f"{relevant_context or 'Нет выделенных фрагментов.'}\n\n"
                            f"Целевой раздел:\n\n"
                            f"{target_context or 'Не выделен.'}\n\n"
                            f"Полный контекст:\n\n{context}\n\n"
                            f"Вопрос: {req.question}"
                        ),
                    },
                ],
                max_tokens=2000,
                temperature=0.5,
            )
        except Exception as e:
            raise HTTPException(502, f"Chat API error: {e}")

        log.info("Raw answer: %r", answer[:200])

        # Убираем блок рассуждений <think>...</think> (deepseek-r1 и подобные)
        answer = re.sub(r"<think>.*?</think>", "", answer, flags=re.DOTALL).strip()
        # Убираем маркеры чанков [CHUNK N] / (CHUNK N), которые LLM может скопировать из контекста
        answer = re.sub(r"\s*[\(\[]\s*CHUNK\s*\d+\s*[\)\]]", "", answer).strip()
        # Нормализуем TeX-команды, которые иногда подсовывает LLM в ответ.
        # Сначала меняем сами команды на юникод-символы, потом снимаем $...$/\(...\)/\[...\] обёртки.
        tex_replacements = [
            (r"\\rightarrow\b", "→"),
            (r"\\leftarrow\b", "←"),
            (r"\\Rightarrow\b", "⇒"),
            (r"\\Leftarrow\b", "⇐"),
            (r"\\leftrightarrow\b", "↔"),
            (r"\\to\b", "→"),
            (r"\\leq\b", "≤"),
            (r"\\le\b", "≤"),
            (r"\\geq\b", "≥"),
            (r"\\ge\b", "≥"),
            (r"\\neq\b", "≠"),
            (r"\\ne\b", "≠"),
            (r"\\approx\b", "≈"),
            (r"\\equiv\b", "≡"),
            (r"\\pm\b", "±"),
            (r"\\times\b", "×"),
            (r"\\cdot\b", "·"),
            (r"\\div\b", "÷"),
            (r"\\infty\b", "∞"),
            (r"\\sum\b", "Σ"),
            (r"\\alpha\b", "α"),
            (r"\\beta\b", "β"),
            (r"\\gamma\b", "γ"),
            (r"\\delta\b", "δ"),
            (r"\\mu\b", "μ"),
            (r"\\pi\b", "π"),
            (r"\\sigma\b", "σ"),
            (r"\\degree\b", "°"),
        ]
        for pattern, repl in tex_replacements:
            answer = re.sub(pattern, repl, answer)
        # Стрелка из дефисов: -> / --> (вне слов вроде e->f)
        answer = re.sub(r"(?<!\w)-+>(?!\w)", "→", answer)
        # Снимаем inline-формулы $...$ / \(...\) / \[...\] — оставляем содержимое
        answer = re.sub(r"\$([^\$\n]{1,80})\$", r"\1", answer)
        answer = re.sub(r"\\\(([^\)\n]{1,80})\\\)", r"\1", answer)
        answer = re.sub(r"\\\[([^\]\n]{1,80})\\\]", r"\1", answer)

        answer_clean = answer.strip()
        for quote in ['"', "'", "«", "»"]:
            answer_clean = answer_clean.strip(quote)
        answer = answer_clean.strip().replace("\\n", "\n")

        # Схлопываем дубли пунктов нумерованного списка и перенумеровываем
        # по порядку (LLM иногда выдаёт ту же фразу под номерами 7..27,
        # если она встречается в контексте в нескольких чанках).
        _list_lines = answer.split("\n")
        _item_re = re.compile(r"^\s*\d+[\.\)]\s+(.*\S)\s*$")
        _seen_items: set[str] = set()
        _new_lines: list[str] = []
        _idx = 0
        for _line in _list_lines:
            _m = _item_re.match(_line)
            if _m:
                _key = re.sub(r"\s+", " ", _m.group(1)).lower()
                if _key in _seen_items:
                    continue
                _seen_items.add(_key)
                _idx += 1
                _new_lines.append(f"{_idx}. {_m.group(1)}")
            else:
                _new_lines.append(_line)
        answer = "\n".join(_new_lines)

        if answer != "Информация в документах не найдена.":
            if not _is_faithful(answer, context):
                log.warning("Faithfulness check failed — suppressing answer")
                answer = "Информация в документах не найдена."

        base = _public_base_url(request)

        # Собираем ссылки на изображения СТРОГО из основного документа ответа.
        image_urls: dict[str, str] = {}
        img_marker_re = re.compile(r"\[Рисунок (\d+): ([^\]]+)\]")

        # Корни слов вопроса (первые 5 букв слов длиной ≥5)
        _q_roots = {w[:5] for w in re.findall(r"[а-яёa-z]{5,}", req.question.lower())}
        _support_scores = _document_support_scores(answer, all_docs, all_metas)

        # Источники: самый релевантный документ (по совпадению имени с вопросом) первым
        _fname_relevance: dict[str, int] = {}
        for fn in _all_sources:
            fn_roots = {w[:5] for w in re.findall(r"[а-яёa-z]{5,}",
                        fn.replace("_", " ").replace("-", " ").lower())}
            _fname_relevance[fn] = len(_q_roots & fn_roots)
        sources = sorted(
            _all_sources,
            key=lambda fn: (_support_scores.get(fn, 0), _fname_relevance.get(fn, 0)),
            reverse=True,
        )
        download_urls = {s: f"{base}/files/{s}" for s in sources}
        _primary_source = sources[0] if sources else None

        _img_chunk_counts: dict[str, int] = {}
        for doc_text, meta in zip(all_docs, all_metas):
            if img_marker_re.search(doc_text):
                fn = meta["filename"]
                _img_chunk_counts[fn] = _img_chunk_counts.get(fn, 0) + 1

        # Источник картинок — primary, если у него они есть.  Иначе —
        # первый по support_score документ, у которого есть картинки
        # в top_k (тогда блок-схема из документа №2 показывается, даже
        # если текст ответа взят из №1).
        if _primary_source and _primary_source in _img_chunk_counts:
            _img_source_file = _primary_source
        else:
            _img_source_file = next(
                (fn for fn in sources if fn in _img_chunk_counts),
                None,
            )

        # Глобальный fallback: ни один документ из top_k не имеет картинок
        # в выдаче (например, картинка лежит только в конце документа,
        # а top_k собрал чанки из начала).  Проверяем по индексу: есть ли
        # у любого из источников вообще image-chunks?  Берём первого
        # такого по support_score.
        if _img_source_file is None and sources:
            for fn in sources:
                try:
                    recs = await _document_records(fn)
                except Exception:
                    continue
                for rec in recs:
                    if img_marker_re.search(rec.get("text", "") or ""):
                        _img_source_file = fn
                        log.info(
                            "[attach_images] fallback to source %s (no img in top_k, but found in index)",
                            fn,
                        )
                        break
                if _img_source_file:
                    break

        log.info(
            "Image source: %s (primary_source=%s, support_scores=%s, img_chunks=%s)",
            _img_source_file,
            _primary_source,
            _support_scores,
            _img_chunk_counts,
        )

        # Собираем чанки и картинки из расширенного контекста (all_docs),
        # но только из файла-источника картинок
        _img_chunks: list[str] = []
        _img_metas: list[dict] = []
        if _img_source_file:
            for doc_text, meta in zip(all_docs, all_metas):
                if meta["filename"] == _img_source_file:
                    _img_chunks.append(doc_text)
                    _img_metas.append(meta)
                    for match in img_marker_re.finditer(doc_text):
                        img_name = match.group(2)
                        marker = match.group(0)
                        if marker not in image_urls:
                            stem = Path(meta["filename"]).stem
                            image_urls[marker] = (
                                f"{base}/files/{stem}_images/{img_name}"
                            )

            # Top_K чанков покрывает только небольшую часть документа.  Чтобы
            # картинки со ВСЕХ страниц документа-источника могли быть
            # привязаны к ответу, подтягиваем все его чанки, содержащие
            # маркеры [Рисунок ...].  Это не влияет на текст, который видит
            # LLM (он уже сформирован), только на пул доступных картинок.
            try:
                src_records = await _document_records(_img_source_file)
            except Exception as e:
                log.warning("[attach_images] failed to load src records: %s", e)
                src_records = []
            seen_chunk_ids = {id(t) for t in _img_chunks}
            extra_image_chunks = 0
            for rec in src_records:
                text = rec.get("text", "") or ""
                if not img_marker_re.search(text):
                    continue
                # дедуп по содержанию, чтобы не задвоить уже взятые чанки
                already = False
                for existing in _img_chunks:
                    if existing == text:
                        already = True
                        break
                if already:
                    continue
                _img_chunks.append(text)
                _img_metas.append({
                    "filename": _img_source_file,
                    "chunk_index": rec.get("chunk_index", -1),
                })
                for match in img_marker_re.finditer(text):
                    img_name = match.group(2)
                    marker = match.group(0)
                    if marker not in image_urls:
                        stem = Path(_img_source_file).stem
                        image_urls[marker] = (
                            f"{base}/files/{stem}_images/{img_name}"
                        )
                extra_image_chunks += 1
            if extra_image_chunks:
                log.info(
                    "[attach_images] pulled %d extra image-chunks from source %s "
                    "(total image_urls now: %d)",
                    extra_image_chunks, _img_source_file, len(image_urls),
                )

        # Привязываем иллюстрации к пунктам ответа — строго из одного документа
        if image_urls and answer != _not_found:
            answer = _attach_images_to_answer(answer, _img_chunks, _img_metas, image_urls)

        answer_html = _answer_to_html(answer, image_urls, download_urls)

        log_entry["answer"] = answer
        log_entry["sources"] = sources
        log_entry["chunks_used"] = len(all_docs)

        return AskResponse(
            answer=answer,
            answer_html=answer_html,
            sources=sources,
            chunks_used=len(all_docs),
            download_urls=download_urls,
            image_urls=image_urls,
            raw_chunks=all_docs[:5],
            rewritten_query=rewritten,
        )
    except HTTPException as e:
        log_entry["error"] = f"HTTP {e.status_code}: {e.detail}"
        raise
    except Exception as e:
        log_entry["error"] = repr(e)
        raise
    finally:
        log_entry["duration_ms"] = int((time.monotonic() - started_mono) * 1000)
        await _log_ask(log_entry)
